from __future__ import annotations

import csv
import json
import shutil
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib import font_manager
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs/result_artifacts/weekly_report_result/report_2026_W34_nb252_optimization"
FIG = OUT / "figures"
TABLE = OUT / "tables"
REPORT = OUT / "Nb252_optimization_stage_report_2026_W31_W34.docx"

BLUE = "2E5D7B"
DARK = "17324D"
MUTED = "66727D"
LIGHT = "EEF3F6"
GOLD = "B78628"
RED = "9B2C2C"
GREEN = "2D6A4F"
INK = "1D252C"

SOURCE_FIGURES = [
    ("input_baseline", ROOT / "docs/result_artifacts/input_baseline/summary/input_baseline_qc.png"),
    ("affinity_scan", ROOT / "docs/result_artifacts/candidate_design/affinity_pyrosetta_full_scan_20260811/affinity_full_scan_qc.png"),
    ("flex_ddg", ROOT / "docs/result_artifacts/candidate_design/flex_ddg_production_result_20260812/flex_ddg_production_qc.png"),
    ("property_landscape", ROOT / "docs/result_artifacts/candidate_design/unified_property_scoring_result_20260815/unified_property_scoring.png"),
    ("double_mutants", ROOT / "docs/result_artifacts/candidate_design/double_mutant_scan_review_v2_1_20260816/double_mutant_joint_evidence_v2_1.png"),
    ("finalist_energy", ROOT / "docs/result_artifacts/candidate_design/finalist_energy_review_20260817/finalist_energy_review.png"),
    ("final_panel", ROOT / "docs/result_artifacts/candidate_design/final_candidate_panel_20260817/final_candidate_panel.png"),
]


def load_json(path: str) -> dict:
    return json.loads((ROOT / path).read_text(encoding="utf-8"))


def load_csv(path: str) -> list[dict[str, str]]:
    with (ROOT / path).open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def set_run_font(run, size=10.5, bold=False, color=INK, name="Calibri", east="Microsoft YaHei"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, header=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "bottom"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "10")
        elem.set(qn("w:color"), DARK)
        borders.append(elem)
    for edge in ("insideH", "insideV", "left", "right"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    if header:
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "7")
            bottom.set(qn("w:color"), "8CA3B2")
            tc_borders.append(bottom)
            tc_pr.append(tc_borders)


def add_table(doc, headers, rows, widths, font_size=8.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(text)
        shade(cell, LIGHT)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, font_size, True, DARK)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(run, font_size, False, INK)
    set_table_geometry(table, widths)
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True, color=DARK)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_callout(doc, label, text, kind="info"):
    colors = {"info": (LIGHT, BLUE), "caution": ("FFF6E5", GOLD), "risk": ("FBECEC", RED), "positive": ("EBF5F0", GREEN)}
    fill, accent = colors[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}  ")
    set_run_font(r, 10.2, True, accent)
    r = p.add_run(text)
    set_run_font(r, 10.2, False, INK)
    p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_caption(doc, text, source):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, 8.5, True, DARK)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)


def add_figure(doc, path, caption, source, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption)
    add_caption(doc, caption, source)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, 8.5, color=MUTED)


def build_workflow_figure():
    cjk_font = Path("C:/Windows/Fonts/msyh.ttc")
    if cjk_font.exists():
        font_manager.fontManager.addfont(str(cjk_font))
        plt.rcParams["font.family"] = "Microsoft YaHei"
        plt.rcParams["axes.unicode_minus"] = False
    stages = [
        ("输入身份基线", "3个结构模型\n47条表达序列"),
        ("亲和力全扫描", "456个单突 × 3重复"),
        ("多构象复核", "48条严格层 + 2条多样性补充"),
        ("统一性质景观", "2318枚举 → 1962评分"),
        ("双突联合分析", "86个组合"),
        ("综合审核", "14单突 + 86双突\n56主池 → 36终审"),
        ("最终面板", "30候选 + 3储备"),
    ]
    fig, ax = plt.subplots(figsize=(11.2, 4.2))
    ax.axis("off")
    xs = [0.07 + i * 0.145 for i in range(len(stages))]
    for i, ((title, note), x) in enumerate(zip(stages, xs)):
        color = "#2E5D7B" if i < len(stages) - 1 else "#2D6A4F"
        ax.text(x, 0.57, title, ha="center", va="center", color="white", fontsize=10, fontweight="bold",
                bbox=dict(boxstyle="round,pad=0.65", fc=color, ec="none"))
        ax.text(x, 0.28, note, ha="center", va="center", fontsize=9, color="#27323A")
        if i < len(stages) - 1:
            ax.annotate("", xy=(xs[i + 1] - 0.055, 0.57), xytext=(x + 0.055, 0.57),
                        arrowprops=dict(arrowstyle="->", color="#8CA3B2", lw=2))
    ax.text(0.5, 0.93, "Nb252多目标计算优化证据漏斗", ha="center", fontsize=16, fontweight="bold", color="#17324D")
    ax.text(0.5, 0.06, "各阶段均保留原始证据、风险标记与实验可检验性；预测结果不等同于实验优化。", ha="center", fontsize=9, color="#66727D")
    fig.tight_layout()
    path = FIG / "figure_02_evidence_funnel.png"
    fig.savefig(path, dpi=600, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    with (TABLE / "evidence_funnel_data.csv").open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(["stage_order", "stage", "display_count"])
        for index, (title, note) in enumerate(stages, 1):
            writer.writerow([index, title, note.replace("\n", "; ")])
    return path


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.167
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("Nb252多目标计算优化 | 阶段周报")
    set_run_font(r, 8.5, True, MUTED)
    add_page_number(section.footer.paragraphs[0])


def write_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("阶段研究周报")
    set_run_font(r, 12, True, GOLD)
    p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NK2R纳米抗体Nb252\n多目标计算优化")
    set_run_font(r, 28, True, DARK)
    p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("从实验结构基线到30条实验测试候选")
    set_run_font(r, 15, False, BLUE)
    p.paragraph_format.space_after = Pt(42)
    add_table(doc, ["报告周期", "项目状态", "报告日期"], [["2026-W31—W34", "计算设计闭环完成", "2026-08-18"]], [3000, 3360, 3000], 9.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目阶段交流材料｜所有候选均为计算优先级，不代表实验验证")
    set_run_font(r, 9, False, MUTED)
    doc.add_page_break()


def write_static_toc(doc):
    add_heading(doc, "目录与阅读导航", 1)
    rows = [
        ("1", "执行摘要与项目结论"), ("2", "科学背景、目标与输入基线"),
        ("3", "结构与界面证据"), ("4", "计算优化总体路线"),
        ("5", "PyRosetta协议与亲和力轨道"), ("6", "表达、稳定性与可开发性工具验证"),
        ("7", "统一单突空间、专家审查与双突组合"), ("8", "终审能量来源与最终30条面板"),
        ("9", "局限性、风险与实验建议"), ("附录", "最终候选、储备/排除项、工具角色与完整序列"),
    ]
    add_table(doc, ["章节", "内容"], rows, [1200, 8160], 9.5)
    add_callout(doc, "阅读建议", "建议先阅读执行摘要与图2证据漏斗，再重点查看第5—8节的筛选逻辑、最终面板和实验建议；附录列出候选概览、关键决策和完整序列。")
    doc.add_page_break()


def build_report():
    OUT.mkdir(parents=True, exist_ok=True)
    FIG.mkdir(exist_ok=True)
    TABLE.mkdir(exist_ok=True)
    for index, (name, source) in enumerate(SOURCE_FIGURES, 1):
        if not source.exists():
            raise FileNotFoundError(source)
        shutil.copy2(source, FIG / f"source_{index:02d}_{name}.png")
    workflow = build_workflow_figure()

    numbering = load_json("docs/run_summaries/input_baseline/sequence_numbering.json")
    alignment = load_json("docs/result_artifacts/input_baseline/structure_released_20260810/structure_alignment_summary.json")
    affinity = load_json("docs/run_summaries/candidate_design/affinity_pyrosetta_full_scan_20260811.json")
    affinity_tiers = load_json("docs/result_artifacts/candidate_design/affinity_post_scan_filter_20260812/affinity_post_scan_gate.json")
    flex = load_json("docs/run_summaries/candidate_design/flex_ddg_production_result_20260812.json")
    affinity_core = load_json("docs/result_artifacts/candidate_design/affinity_ensemble_core_20260813/affinity_ensemble_core_gate.json")
    unified_plan = load_json("docs/result_artifacts/candidate_design/unified_single_mutant_plan_20260815/unified_single_mutant_plan_gate.json")
    properties = load_json("docs/run_summaries/candidate_design/unified_property_scoring_result_20260815.json")
    property_review = load_json("docs/result_artifacts/candidate_design/unified_property_scoring_result_20260815/unified_property_scoring_scientific_review.json")
    property_affinity = load_json("docs/result_artifacts/candidate_design/property_affinity_pyrosetta_review_20260816/property_affinity_scientific_review.json")
    single_shortlist = load_json("docs/result_artifacts/candidate_design/single_mutant_shortlist_20260816/single_mutant_shortlist_gate.json")
    doubles = load_json("docs/run_summaries/candidate_design/double_mutant_scan_review_v2_1_20260816.json")
    preliminary = load_json("docs/run_summaries/candidate_design/preliminary_panel_20260817.json")
    energy_gate = load_json("docs/result_artifacts/candidate_design/finalist_energy_review_20260817/finalist_energy_review_gate.json")
    final_gate = load_json("docs/result_artifacts/candidate_design/final_candidate_panel_20260817/final_candidate_panel_gate.json")
    candidates = load_csv("docs/result_artifacts/candidate_design/final_candidate_panel_20260817/final_candidates_30.csv")
    reserve_audit = load_csv("docs/result_artifacts/candidate_design/final_candidate_panel_20260817/final_reserve_status.csv")
    decision_audit = load_csv("docs/result_artifacts/candidate_design/final_candidate_panel_20260817/final_candidate_decision_audit.csv")

    doc = Document()
    configure_document(doc)
    write_cover(doc)
    write_static_toc(doc)

    add_heading(doc, "1. 执行摘要与项目结论", 1)
    add_callout(doc, "阶段结论", "项目已完成结构身份冻结、亲和力单突扫描、多构象复核、统一性质评价、双突组合、结构/化学风险审查和显式终审，形成30条互不重复的实验测试候选，另保留3条储备。", "positive")
    add_body(doc, "本项目以实验解析的NK2R–Nb252复合物为结合构象主证据，在保持当前表位、构象和末端SSGS不变的前提下，同时探索亲和力、稳定性、表达相关性质与可开发性。计算流程没有把单一模型作为真值，而是将PyRosetta、Flex ddG、AntiFold、NetSolP、NanoMelt、TNP和序列语言模型证据分层使用。")
    add_body(doc, "最终30条面板由6条亲和力导向单突、6条性质导向单突、17条平衡组合和1条亲和力支持双突组成。WT作为独立实验对照，不占30条候选名额。所有序列均保持128 aa、Cys22/Cys95及末端SSGS。")
    add_callout(doc, "解释边界", "最终面板的状态是“建议实验测试”，不是“已获得30条优化抗体”。计算能量、Tm、溶解性、序列兼容性和表面风险均不能替代实测KD、表达量、SEC单体比例或热稳定性。", "caution")
    add_table(doc, ["阶段指标", "结果", "状态"], [
        ["表达序列审核", f"{numbering['acceptance']['numbering_pass_count']}成功 / {numbering['acceptance']['numbering_failed_count']}失败", "完成"],
        ["亲和力全扫描", f"{affinity['counts']['candidate_count']}候选 × {affinity['counts']['replicate_count_per_candidate']}重复", "完成"],
        ["Flex ddG", f"{flex['candidate_count']}候选 / {flex['task_count']}任务", "完成"],
        ["统一性质单突", f"{properties['counts']['candidates']}候选 + {properties['counts']['wt_controls']} WT", "完成"],
        ["双突联合分析", f"{doubles['candidate_count']}候选", "完成"],
        ["最终面板", f"{final_gate['final_candidate_count']}候选 + {final_gate['reserve_count']}储备", "建议进入实验测试"],
    ], [2800, 3700, 2860], 8.8)

    add_heading(doc, "2. 科学背景、目标与输入基线", 1)
    add_body(doc, "Nb252是针对NK2R的实验筛选纳米抗体。本项目的目标不是仅追求某一个计算分数，而是在维持实验观察表位与结合构象的基础上，提高亲和力，同时避免明显损害稳定性、表达和可开发性，并给出具有机制多样性的实验面板。")
    add_heading(doc, "2.1 三类结构证据", 2)
    add_bullet(doc, "NK2R-252.pdb：实验解析的NK2R–Nb252结合构象，是界面和表位的主要结构证据。")
    add_bullet(doc, "NK2R-NKA：提供NKA结合背景，不直接证明Nb252突变效应或功能机制。")
    add_bullet(doc, "fold_2r_252_nomg_model_0.cif：AF3预测的Nb252结构，只用于补充未解析区域和独立构象参考。")
    add_heading(doc, "2.2 序列与实验约束", 2)
    add_body(doc, "合作者确认完整128-aa reported Nb252为权威设计母本；末端SSGS不是linker，候选设计时不得替换或删除。47条产量序列来自BL21体系，且不同来源数据可直接比较；但LLJ记录保留分组/删失语义，不将其虚构为个体精确点估计。")
    add_body(doc, f"ANARCII provisional IMGT审核中，{numbering['acceptance']['numbering_pass_count']}条成功，WCC__4-28因模型分数低于识别阈值而失败。Nb252前126 aa被编号，末端GS未获IMGT编号，但仍属于权威构建体并受SSGS冻结规则保护。")
    add_figure(doc, FIG / "source_01_input_baseline.png", "图1  Nb252输入身份、编号与覆盖基线", "input_baseline/summary/input_baseline_qc.png")

    add_heading(doc, "3. 结构与界面证据", 1)
    add_body(doc, "ChimeraX会话中的三个模型已分别导出为可追溯mmCIF，并由Gemmi读回。链角色经人工确认：实验复合物C链为Nb252、R链为NK2R；NKA模型L链为NKA、R链为NK2R；AF3模型A链为Nb252。")
    add_body(doc, "实验VHH存在未解析坐标，关键缺失位置已经固化为reported indices 9–15、24–29；这些位置在实验界面分析中保持not_evaluable，而不是被解释为非界面。")
    add_table(doc, ["比较项", "结果", "解释"], [
        ["框架拟合", "82个共同framework Cα", "CDR3不参与框架拟合"],
        ["框架RMSD", "0.631994 Å", "实验与AF3框架总体一致"],
        ["CDR3 RMSD", "6.490853 Å", "CDR3构象差异显著"],
        ["CDR3最大位移", "10.733119 Å", "AF3不作为实验CDR3真值"],
    ], [2300, 2500, 4560], 8.7)
    add_body(doc, "界面采用同一实验复合物内VHH与NK2R聚合物重原子的严格中心距离<4.0 Å定义，排除氢/氘、非正occupancy、水、配体、糖、离子、晶体/NCS副本，并遵守altloc兼容性。重算得到24个VHH界面残基、246对原子接触，与合作者橙色标注完全一致。界面位点是“谨慎突变区”，不是绝对禁区或能量热点清单。")
    add_callout(doc, "AF3使用原则", "由于WT的AF3 CDR3与实验结合构象偏差较大，AF3未用于常规亲和力排序；它只作为缺失区和独立局部包装敏感性证据。", "caution")

    add_heading(doc, "4. 计算优化总体路线", 1)
    add_figure(doc, workflow, "图2  从输入基线到最终30条序列的证据漏斗", "基于各阶段计算与审查结果汇总")
    add_body(doc, "路线采用“宽空间计算、统一证据筛选、显式专家终审”。亲和力与性质并非严格按CDR/FR分区：界面和非界面均可提出候选，但每个突变都必须同时接受结构、亲和力、性质和化学风险审查。")

    add_heading(doc, "5. PyRosetta协议与亲和力轨道", 1)
    add_heading(doc, "5.1 评分协议校准", 2)
    add_body(doc, "实验结构直接使用ref2015时总分较高，主要源于输入侧链缺原子、局部rotamer与Rosetta统计势不完全匹配以及未优化的界面排斥。项目没有把该绝对分数用于候选结论，而是校准并固定interface_repack_constrained_min协议，以同一prepared WT、相同局部约束和配对WT为参照比较突变体。")
    add_body(doc, "interface_repack_constrained_min是本项目对局部结构准备流程的命名：先只在界面邻域重新选择侧链rotamer，再对同一邻域的主链和侧链做坐标约束下的有限能量最小化；其余残基及链间刚体位置不动。它用于缓解实验模型中的局部碰撞，同时尽量维持已解析的结合构象，不是全局relax或重新对接。WT与突变体使用相同的可动范围、约束和评分设置，因此比较的是同一实验构象附近的配对相对能量。")
    add_body(doc, "所选协议在8次重复中得到稳定的负dG_separated和跨界面能，并把界面fa_rep相对raw结构显著降低，同时保持很小的界面Cα位移。由于本项目关注同一实验构象附近的相对排序，未切换RosettaMP、未全局relax，也未花费大量计算补全远离已解析界面的缺失主链。")
    add_callout(doc, "能量解释", "Rosetta ΔdG、跨界面能和total score均为模型内排序信号。尤其是负ΔdG可能由分离态被惩罚而产生，不能直接等价于实验亲和力提高。", "risk")
    add_heading(doc, "5.2 456条单突全扫描", 2)
    add_body(doc, f"界面24个位点形成456条氨基酸替换候选，每条运行3次配对重复，共{affinity['counts']['mutant_evaluation_count']}次突变体评价。扫描阶段不设置提前筛选，456条均完成并通过运行完整性门。")
    add_figure(doc, FIG / "source_02_affinity_scan.png", "图3  456条PyRosetta亲和力单突全扫描质量控制", "affinity_pyrosetta_full_scan_20260811")
    tier_counts = affinity_tiers["tier_counts"]
    add_body(doc, "全扫描完成后才统一分层。分层同时考虑两项能量方向、三次重复的一致性、接触保持、局部排斥和界面骨架稳定性，而不是按单一最低分截取。")
    add_table(doc, ["层级", "数量", "进入该层的主要依据", "后续用途"], [
        ["Tier 1", tier_counts["tier_1"], "3/3重复中ΔdG与跨界面能均有利；VHH/NK2R接触完全保持；局部fa_rep不升高", "严格复核池"],
        ["Tier 2", tier_counts["tier_2"], "3/3重复中两项能量均有利且受体表位接触保持，但局部包装或VHH接触仍需复核", "严格复核池"],
        ["Tier 3", tier_counts["tier_3"], "两项能量中位数均有利，但重复一致性弱于Tier 1/2", "仅作补充来源"],
        ["Tier 4", tier_counts["tier_4"], "ΔdG与跨界面能方向不一致", "不进入本轮Flex ddG"],
        ["Tier 5", tier_counts["tier_5"], "两项能量中位数未同时有利", "不进入本轮Flex ddG"],
    ], [950, 850, 4900, 2660], 7.8)
    add_heading(doc, "5.3 Flex ddG多构象复核", 2)
    add_body(doc, f"50条复核对象由48条严格复核层候选——全部Tier 1/2（{tier_counts['tier_1']}+{tier_counts['tier_2']}=48条）——和两条位置多样性Tier 3补充候选D33N、Y115F组成。后两条用于避免复核池只覆盖最强能量位点，并不代表它们优于其余Tier 3；其余37条Tier 3本轮不进入高成本复核。")
    add_body(doc, f"每条候选运行20个独立Backrub样本，共{flex['task_count']}个任务。Backrub在突变位点周围8 Å局部主链邻域采样构象，并在同一主链样本上成对比较WT与突变体，从而检验三重复单一prepared结构上的有利方向能否在构象扰动下保持。")
    add_figure(doc, FIG / "source_03_flex_ddg.png", "图4  50条候选Flex ddG多构象复核", "flex_ddg_production_result_20260812")
    add_heading(doc, "5.4 从50条Flex ddG候选到亲和力单突", 2)
    add_body(doc, f"亲和力核心门要求ΔdG和跨界面能各自在至少18/20个样本中为负，且两项中位数均为负。50条中共有{affinity_core['core_module_count']}条通过，分布在{affinity_core['core_position_count']}个位点；这是多构象方向一致性证据，不是实测亲和力。")
    add_body(doc, "随后进行结构与化学专家审查：新增未配对Cys的R45C被直接移出，较低风险且具补充支持的R45T作为同位点替代假设加入。由此形成8条活跃亲和力单突：R45T、R45V、D101W、I103W、E105F、E105L、N107A和S114M。终审又因D101W和I103W的突变体复合物在三次配对重复中均不利、且同时具有暴露Trp/接触/局部排斥风险而排除，最终保留6条亲和力导向单突。")

    add_heading(doc, "6. 表达、稳定性与可开发性工具验证", 1)
    add_heading(doc, "6.1 工具与真实产量的关系", 2)
    add_body(doc, "项目首先在47条BL21产量序列上验证可用工具与真实产量之间的关系。nanoBERT、NetSolP、TNP和NanoMelt均未提供足够稳健、可迁移的产量排序能力，因此没有训练47样本局部模型，也没有用任一预测器替代实验产量。")
    add_table(doc, ["工具", "评价对象", "项目内用途", "证据限制"], [
        ["nanoBERT", "抗体序列语言合理性", "背景关联验证", "不支持候选产量排序"],
        ["NetSolP", "Solubility / Usability", "兼容性与极端风险过滤", "变化幅度通常较小"],
        ["NanoMelt", "预测表观Tm", "稳定性辅助排序", "不是实测Tm"],
        ["TNP", "PSH/PPC/PNC与flag", "表面疏水/电荷风险", "flag只作辅助风险证据"],
        ["AntiFold", "结构条件下序列兼容性", "三结构视图一致性", "不单独决定候选"],
        ["PyRosetta", "复合物/分离态相对能量", "亲和力及局部结构风险", "REU不能跨协议直接比较"],
    ], [1450, 2100, 2700, 3110], 8.2)
    add_body(doc, "本报告后续使用的四项核心性质指标均以候选相对WT的变化解释；对以下四项而言，正向变化表示模型预测更有利，但不等于实验改善。")
    add_table(doc, ["指标", "简要含义", "本项目中的方向与边界"], [
        ["NetSolP U", "NetSolP给出的Usability（综合可用性）0–1模型分数", "越高越有利；仅作可开发性相容信号，不等于BL21产量"],
        ["NetSolP S", "NetSolP给出的Solubility（溶解性）0–1模型分数", "越高表示预测溶解性倾向越好；不等于实测溶解度或回收率"],
        ["NanoMelt预测Tm", "由VHH序列预测的表观熔解温度，单位为°C", "越高表示预测热稳定性越好；属于计算预测，不是DSF等实验Tm"],
        ["实验复合物视图AntiFold", "在实验NK2R–Nb252复合物结构背景下，比较突变氨基酸与WT的条件对数概率", "ΔlogP>0表示突变与该固定结构背景更兼容；不是结合自由能、亲和力或表达量"],
    ], [1900, 3500, 5560], 8.0)
    add_heading(doc, "6.2 性质候选如何从全空间选出", 2)
    add_body(doc, f"性质路线首先在122个可变reported位置枚举19种非WT替换，共{unified_plan['candidate_count']}条。随后一次性应用硬约束：{unified_plan['design_status_counts']['blocked_new_unpaired_cys']}条因新增未配对Cys阻断，{unified_plan['design_status_counts']['deferred_missing_experimental_coordinates']}条因实验结构缺失坐标而推迟，最终{unified_plan['design_status_counts']['eligible_current_round']}条进入统一评分。该评分集包括432条已存在的非Cys界面候选和1530条非界面性质发现候选。")
    property_track = property_review["track_review"]["stability_developability_discovery"]
    add_body(doc, f"在1530条性质发现候选中，以NetSolP U、NetSolP S、NanoMelt预测Tm和实验复合物视图AntiFold兼容性进行同轨道Pareto分层，得到{property_track['pareto_front_1_count']}条Pareto 1。Pareto 1仅表示没有另一条候选在全部指标上同时更优，并不表示四项指标都改善。")
    add_table(doc, ["性质筛选阶段", "候选数", "筛选规则或作用"], [
        ["全序列枚举", unified_plan["candidate_count"], "122个可变位置 × 19种非WT替换"],
        ["统一可评分集合", unified_plan["design_status_counts"]["eligible_current_round"], "排除新增未配对Cys；缺失坐标候选本轮暂缓"],
        ["性质发现Pareto 1", property_track["pareto_front_1_count"], "四项性质信号的非支配集合；不直接等于入选"],
        ["幅度感知复核池", property_affinity["validation"]["candidate_count"], "至少1项U/S/Tm达到非微小改善；无非微小恶化；无新增化学风险"],
        ["PyRosetta方向有利", property_affinity["validation"]["direction_class_counts"]["directionally_favorable"], "两项能量中位数均有利，且各至少2/3重复有利"],
        ["最终活跃性质单突", single_shortlist["active_after_by_track"]["property"], "综合亲和力、接触、AntiFold、缺失区局部包装和化学风险"],
    ], [2500, 1050, 5810], 8.0)
    add_body(doc, "49条Pareto 1先接受TNP表面风险复核；TNP没有作为加分器或硬排序器。幅度感知门随后选出30条、覆盖10个位点的性质复核池：必须至少有一项U/S/Tm达到预设的非微小改善，同时三项均不得出现非微小恶化，且不得新增糖基化、脱酰胺、异构化或额外M/W等化学风险。AntiFold与TNP保留为独立证据列，不参与这一30条的硬门。")
    add_body(doc, f"30条再进行3次配对PyRosetta亲和力非劣复核，结果为{property_affinity['validation']['direction_class_counts']['directionally_favorable']}条方向有利、{property_affinity['validation']['direction_class_counts']['mixed']}条混合、{property_affinity['validation']['direction_class_counts']['directionally_adverse']}条方向不利。结合接触保持、暴露疏水/强负AntiFold、实验缺失区邻近位点的AF3局部包装及不可补偿Pro风险，先形成22条活跃性质假设，再剔除16条负向证据更强的候选，保留Q1D、A23S、F30A、F30S、F30T和S55G共6条性质单突。")
    add_body(doc, "被剔除的16条主要原因可归为：4条受体接触变化、3条AF3局部非劣门未通过、2条配对亲和力方向不利、3条强负AntiFold并伴随暴露疏水风险，以及4条其他强负AntiFold复合物信号。各原因按主阻断项计数，避免同一候选重复统计。")
    add_figure(doc, FIG / "source_04_property_landscape.png", "图5  统一单突性质景观与分层结果", "unified_property_scoring_result_20260815")

    add_heading(doc, "7. 统一单突空间、专家审查与双突组合", 1)
    add_heading(doc, "7.1 统一候选合同", 2)
    add_body(doc, "设计合同不再冻结整个框架区，也不把CDR和FR简单划分为亲和力区与性质区。亲和力候选可容许性质指标轻微负向，但阻止明确且一致的恶化；性质候选至少需要一个非微小改善，且不能出现明显的亲和力或结构破坏。")
    add_body(doc, "经过前述多构象、性质幅度和专家安全门后，活跃单突固定为14条：8条亲和力来源和6条性质来源。它们既是可独立测试的假设，也是解释双突组合效应所需的组成单突对照。")
    add_heading(doc, "7.2 专家化风险审查", 2)
    add_body(doc, "所有核心单突从结构和可开发性角度复核异常二硫键、新生糖基化motif、暴露疏水替换、致密疏水窗口、氧化风险、CDR柔性变化、保守位点、局部fa_rep、接触改变和对结构准备的敏感性。F30P等具有难以弥补的骨架风险突变直接移出候选空间，避免重复计算低价值假设。")
    add_heading(doc, "7.3 86条双突联合分析", 2)
    jc = doubles["joint_evidence_class_counts"]
    add_body(doc, "双突并非只将一个亲和力突变与一个性质突变组合，而是对14条活跃单突进行完整两两组合。理论上14选2为91组；去除同一位置不能同时发生的5组互斥替换（R45的1组、E105的1组、F30的3组）后，得到86条可实现双突，并全部在统一协议下重新评分。")
    add_table(doc, ["联合证据类别", "数量", "含义"], [
        ["亲和力与性质共同支持", jc["balanced_supported"], "优先的平衡候选来源"],
        ["亲和力支持、性质非不利", jc["affinity_supported_property_nonadverse"], "偏亲和力候选"],
        ["性质支持、亲和力非不利", jc["property_supported_affinity_nonadverse"], "偏性质候选"],
        ["权衡或无清晰联合支持", jc["tradeoff_or_no_clear_joint_support"], "保留审计但不直接优先"],
    ], [3900, 1300, 4160], 8.6)
    add_body(doc, f"86/86均通过PyRosetta结构安全门；{doubles['paired_contact_audit']['paired_contact_changed_candidate_count']}条相对配对WT发生接触集合变化，{doubles['paired_contact_audit']['paired_contact_unchanged_candidate_count']}条不变。接触集合并不要求逐原子精确一致，而是结合表位保持、变化规模和能量方向判断。{doubles['experimental_reference_sensitivity_status_counts']['sensitive']}条对实验参考定义敏感，{doubles['experimental_reference_sensitivity_status_counts']['not_sensitive']}条相对稳健。")
    add_figure(doc, FIG / "source_05_double_mutants.png", "图6  86条双突的亲和力—性质联合证据", "double_mutant_scan_review_v2_1_20260816")

    add_heading(doc, "8. 终审能量来源与最终30条面板", 1)
    add_heading(doc, "8.1 从100条审查记录到36条终审候选", 2)
    add_body(doc, f"这里的100条不是新生成的额外序列，而是前一步已经固定的14条活跃单突加86条可实现双突。14条单突全部进入主证据池，因为它们既是独立候选，也是解释组合所必需的组成对照；86条双突中，只有联合证据属于“亲和力与性质共同支持”“亲和力支持且性质非不利”或“性质支持且亲和力非不利”，并同时通过结构安全与硬约束的42条进入主证据池。因而主证据池为14+42={preliminary['primary_pool_count']}条，其余44条权衡或证据不清晰双突仅保留审计。")
    add_body(doc, "42条支持双突分为28条平衡支持、12条亲和力支持且性质非不利、2条性质支持且亲和力非不利。初步面板保留全部14条组成单突，并从28条平衡支持双突中选择16条；选择顺序先比较同一双突协议内的四项目标Pareto层——ΔdG、跨界面能、非微小性质改善项数和AntiFold兼容性——再用突变与位置对多样性打破同层排序。任何组成突变最多进入5条、任何位置对最多进入2条，防止R45或30/45等单一假设垄断面板。")
    add_body(doc, f"这样形成{preliminary['preliminary_panel_count']}条初步面板（14条单突+16条平衡双突）。同时按三类联合证据各保留2条储备：2条平衡双突、2条亲和力支持双突和2条性质支持双突，共{preliminary['reserve_count']}条。初步30条与6条储备合并为36条终审候选；随后复用既有三重复PyRosetta结果进行能量来源分解，没有新增PyRosetta或AF3计算。")
    add_table(doc, ["漏斗阶段", "数量", "组成与去向"], [
        ["统一审查全集", preliminary["reviewed_candidate_count"], "14条活跃单突 + 86条双突"],
        ["主证据池", preliminary["primary_pool_count"], "14条单突 + 42条支持双突；44条权衡/不清晰双突留档"],
        ["初步面板", preliminary["preliminary_panel_count"], "全部14条单突 + 16条多样化平衡双突"],
        ["分层储备", preliminary["reserve_count"], "平衡/亲和力支持/性质支持双突各2条"],
        ["能量来源终审", energy_gate["candidate_count"], "初步30条 + 6条储备，逐条显式决定"],
    ], [2300, 950, 6110], 8.2)
    ec = energy_gate["energy_origin_class_counts"]
    add_table(doc, ["能量来源类别", "数量", "终审解释"], [
        ["复合物与分离态均有利", ec["complex_and_separated_state_stabilization"], "相对更直接的双态支持"],
        ["表观结合改善由分离态惩罚驱动", ec["apparent_binding_gain_driven_by_separated_destabilization"], "不得单独声称亲和力提高"],
        ["一致分离态谨慎", ec["consistent_separated_destabilization_caution"], "结合其他证据审查"],
    ], [4550, 1300, 3510], 8.5)
    add_figure(doc, FIG / "source_06_finalist_energy.png", "图7  36条终审候选的配对Rosetta能量来源审查", "终审能量来源结果")
    add_heading(doc, "8.2 最终候选确定", 2)
    add_body(doc, "终审逐条结合直接复合物能量、跨界面能、分离态来源、接触/表位保持、AntiFold、U/S/Tm、TNP、化学风险、组成单突对照价值和面板多样性。最终排除D101W、I103W和Q1D;R45V；保留F30S;E105L、R45T;E105L、S55G;E105F为储备；从原储备提升A23S;R45V、F30T;E105F、R45T;E105F。")
    cat = final_gate["final_category_counts"]
    add_table(doc, ["最终类别", "数量", "实验作用"], [
        ["亲和力导向单突", cat["affinity_focused_single"], "检验单点界面假设与组合组成效应"],
        ["性质导向单突", cat["property_focused_single"], "检验稳定性/可开发性假设"],
        ["平衡组合", cat["balanced_combination"], "联合亲和力与性质证据"],
        ["亲和力支持双突", cat["affinity_supported_double"], "补充直接亲和力组合假设"],
    ], [3300, 1200, 4860], 8.7)
    add_figure(doc, FIG / "source_07_final_panel.png", "图8  最终30条候选的类别、能量来源与风险构成", "final_candidate_panel_20260817")
    add_callout(doc, "终审结论", "最终确定30条互不重复的建议测试序列，另保留3条储备并排除3条风险或证据不足的候选。所有候选均为计算优先结果，尚需实验验证。", "positive")

    add_heading(doc, "9. 局限性、风险与实验建议", 1)
    add_heading(doc, "9.1 当前局限", 2)
    add_bullet(doc, "尚无候选的实测亲和力、表达量、稳定性、SEC或功能数据。")
    add_bullet(doc, "AF3对WT CDR3的构象偏差较大，不能用作结合构象真值。")
    add_bullet(doc, "部分负ΔdG由分离态惩罚驱动；最终面板仍保留少量低置信机制/性质假设以维持实验信息量。")
    add_bullet(doc, "NetSolP、NanoMelt、TNP、nanoBERT与47条BL21产量的关系较弱，均不具备实测产量替代能力。")
    add_bullet(doc, "双突效应可能存在非加和性；计算排序只能确定测试优先级。")
    add_heading(doc, "9.2 推荐实验设计", 2)
    add_body(doc, "建议将30条候选与WT放在同一构建体、表达和纯化流程中测试。第一轮优先采用小规模BL21表达和统一亲和力测定，以最小成本识别明确失败和有价值信号；第二轮再对表现优良候选开展更完整的稳定性、聚集和功能表征。")
    add_table(doc, ["实验模块", "主要读出", "关键对照/判定目的"], [
        ["BL21表达", "总表达、可溶比例、纯化回收率", "WT同批；验证真实产量"],
        ["纯度与聚集", "SDS-PAGE、SEC单体比例", "排除聚集与异常构象"],
        ["稳定性", "DSF/DSC Tm、应力后单体比例", "验证NanoMelt等预测方向"],
        ["结合", "KD、kon、koff或统一替代读出", "WT与单突组成对照；验证亲和力"],
        ["表位/功能", "竞争、细胞或功能实验", "确认表位与作用模式保持"],
    ], [1700, 3300, 4360], 8.5)
    add_callout(doc, "预定义决策原则", "优先保留亲和力改善且表达/单体比例不明显下降的候选；性质改善候选必须证明不损害结合。所有阴性、近中性与失败结果也应保留，用于反向评估计算路线。")

    doc.add_page_break()
    add_heading(doc, "附录A  最终30条候选概览", 1)
    category_cn = {
        "affinity_focused_single": "亲和力单突", "property_focused_single": "性质单突",
        "balanced_combination": "平衡组合", "affinity_supported_double": "亲和力双突",
    }
    risk_cn = {
        "low": "低", "medium": "中", "medium_high": "中高", "high": "高",
        "no_contact_change": "未见接触变化", "contact_review": "需关注局部接触",
    }
    energy_cn = {
        "complex_and_separated_state_stabilization": "复合物与分离态均有利",
        "apparent_binding_gain_driven_by_separated_destabilization": "表观结合改善主要由分离态惩罚驱动",
        "consistent_separated_destabilization_caution": "分离态持续不利，需谨慎解释",
    }
    category_role = {
        "affinity_focused_single": "检验单点亲和力假设",
        "property_focused_single": "检验性质改善及组合组成效应",
        "balanced_combination": "检验亲和力与性质的联合改善",
        "affinity_supported_double": "检验双位点亲和力协同",
    }
    rows = []
    for row in candidates:
        favorable = int(row["property_material_favorable_count"])
        adverse = int(row["property_material_adverse_count"])
        contact = "局部接触保持" if row["pyrosetta_contact_change_status"] in {"unchanged", "reviewed_in_single_mutant_source"} else "存在已审查的局部接触变化"
        reason = f"{energy_cn[row['energy_origin_class']]}；{contact}；性质非微小改善{favorable}项、明显不利{adverse}项；{category_role[row['panel_category']]}。"
        rows.append([
            row["final_panel_order"], row["mutation_set"], category_cn[row["panel_category"]],
            risk_cn[row["expert_risk_level"]], energy_cn[row["energy_origin_class"]], reason,
        ])
    add_table(doc, ["序号", "突变", "类别", "风险", "能量来源", "入选要点"], rows, [600, 1200, 1300, 800, 2200, 3260], 6.9)
    add_body(doc, "说明：风险等级和能量来源均为计算审查标签，不是实验表型。完整128-aa序列见附录D；WT为面板外独立对照。")

    add_heading(doc, "附录B  储备、排除与关键决策", 1)
    decision_cn = {"reserve": "储备", "exclude": "排除"}
    external_decision_reason = {
        "F30S;E105L": "复合物能量和AntiFold方向有利，但性质变化未达到非微小改善标准，且E105L已达到预设多样性上限。",
        "R45T;E105L": "复合物和结合能量证据较强，但有一项性质指标明显不利，且E105L已达到预设多样性上限。",
        "S55G;E105F": "复合物及跨界面能量改善，但有一项性质指标明显不利，且E105F已有更强的互补组合。",
        "D101W": "三次配对重复中突变体复合物均不利，表观结合改善完全由分离态惩罚驱动，并伴随暴露Trp、氧化、接触变化和局部排斥风险。",
        "I103W": "三次配对重复中突变体复合物均不利，表观结合改善完全由分离态惩罚驱动，并伴随Trp氧化、接触变化、局部排斥和AntiFold负向证据。",
        "Q1D;R45V": "突变体复合物明显不利，表观结合改善完全由分离态惩罚驱动；A23S;R45V提供了直接证据更好的同位点替代方案。",
    }
    reserve_rows = []
    for row in reserve_audit:
        decision = row.get("review_decision", row.get("final_status", ""))
        if decision in {"reserve", "exclude"}:
            mutation = row.get("mutation_set", row.get("candidate_id", ""))
            reserve_rows.append([mutation, decision_cn[decision], external_decision_reason[mutation]])
    for row in decision_audit:
        if row.get("review_decision") == "exclude":
            reserve_rows.append([row["mutation_set"], "排除", external_decision_reason[row["mutation_set"]]])
    add_table(doc, ["候选", "决定", "核心理由"], reserve_rows, [2000, 1200, 6160], 8.2)

    add_heading(doc, "附录C  软件与证据角色", 1)
    add_table(doc, ["工具", "在本项目中的作用", "不能据此声称"], [
        ["PyRosetta / Flex ddG", "同一实验构象附近的配对相对能量、局部构象与重复稳健性", "实验KD、kon、koff或真实自由能"],
        ["AntiFold", "在既定结构背景下评价序列兼容性", "表达量或热稳定性"],
        ["NetSolP", "溶解性与usability的辅助方向信号", "BL21实测产量"],
        ["NanoMelt", "相对WT的预测表观Tm方向", "实测Tm"],
        ["TNP", "表面疏水/电荷与可开发性风险提示", "产量排序或单独淘汰依据"],
        ["nanoBERT", "验证语言模型分数与现有产量数据的关系", "候选产量预测"],
        ["AF3", "补充实验缺失区与条件性局部包装复核", "实验CDR3构象或亲和力排序"],
    ], [1900, 4300, 3160], 8.2)

    add_heading(doc, "附录D  最终30条完整序列", 1)
    add_body(doc, "以下序列按最终面板顺序列出。每条均为128 aa并保留末端SSGS；这些序列属于建议实验测试候选，尚未经实验验证。")
    for row in candidates:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(f"{row['final_panel_order']}. {row['mutation_set']}")
        set_run_font(r, 8.2, True, DARK)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(row["sequence"])
        set_run_font(r, 7.2, False, INK, name="Consolas", east="Microsoft YaHei")

    sources = [
        "docs/run_summaries/input_baseline/sequence_numbering.json",
        "docs/result_artifacts/input_baseline/structure_released_20260810/structure_alignment_summary.json",
        "docs/result_artifacts/input_baseline/interface_released_20260810/interface_manifest.json",
        "docs/run_summaries/candidate_design/affinity_pyrosetta_full_scan_20260811.json",
        "docs/run_summaries/candidate_design/flex_ddg_production_result_20260812.json",
        "docs/run_summaries/candidate_design/unified_property_scoring_result_20260815.json",
        "docs/run_summaries/candidate_design/double_mutant_scan_review_v2_1_20260816.json",
        "docs/result_artifacts/candidate_design/finalist_energy_review_20260817/finalist_energy_review_gate.json",
        "docs/result_artifacts/candidate_design/final_candidate_panel_20260817/final_candidate_panel_gate.json",
        "docs/result_artifacts/candidate_design/final_candidate_panel_20260817/final_candidates_30.csv",
        "docs/result_artifacts/candidate_design/final_candidate_panel_20260817/final_candidates_30.fasta",
        "docs/history/2026-W31.md; 2026-W32.md; 2026-W33.md; 2026-W34.md",
    ]

    doc.core_properties.title = "NK2R纳米抗体Nb252多目标计算优化阶段周报"
    doc.core_properties.subject = "2026-W31—W34项目阶段总结"
    doc.core_properties.author = "Antibody_optimization project"
    doc.core_properties.keywords = "Nb252, NK2R, nanobody, affinity, stability, expression"
    doc.save(REPORT)

    manifest = {
        "schema_version": 1,
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "report": str(REPORT.relative_to(ROOT)).replace("\\", "/"),
        "reporting_period": "2026-W31--W34",
        "audience": "collaborator_and_supervisor",
        "internal_path_index_in_report": False,
        "final_candidate_count": len(candidates),
        "source_figures": [str(path.relative_to(ROOT)).replace("\\", "/") for _, path in SOURCE_FIGURES],
        "source_records": sources,
        "design_preset": "standard_business_brief",
        "header_template": "editorial_cover",
        "table_override": "three-line scientific tables required by AGENTS.md",
    }
    (OUT / "report_manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    (OUT / "README.md").write_text(
        "# Nb252阶段周报制品\n\n"
        "本目录保存面向合作者和导师的2026-W31—W34累计阶段周报、报告专用证据漏斗图、精确漏斗数据和复用图。"
        "正文解释亲和力、性质及终审候选的逐级筛选过程，不展示项目内部相对路径索引。"
        "周报中的数值来自`report_manifest.json`列出的权威制品；所有候选均为计算优先级，尚未实验验证。\n",
        encoding="utf-8",
    )
    return REPORT


if __name__ == "__main__":
    print(build_report())
