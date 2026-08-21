from __future__ import annotations

import csv
import json
import shutil
from collections import Counter
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib import font_manager
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs/result_artifacts/weekly_report_result/report_2026_W34_nb252_expression_route"
FIG = OUT / "figures"
TABLE = OUT / "tables"
DELIVERY = OUT / "delivery"

SELECTION_DIR = ROOT / "docs/result_artifacts/candidate_design/expression_single_mutant_trial_selection_v2_20260820"
CONSERVATION_DIR = ROOT / "docs/result_artifacts/input_baseline/vhh_conservation_consensus_v2_20260819"
LOGO_DIR = ROOT / "docs/result_artifacts/input_baseline/vhh_sequence_logos_20260818"
LANDSCAPE_DIR = ROOT / "docs/result_artifacts/candidate_design/expression_single_mutant_landscape_v1_20260820"
NETSOLP_RESULT_DIR = ROOT / "docs/result_artifacts/candidate_design/netsolp_yield_validation_result_20260814"
NANOMELT_RESULT_DIR = ROOT / "docs/result_artifacts/candidate_design/nanomelt_yield_validation_result_20260815"
NANOMELT_CLASS_DIR = ROOT / "docs/result_artifacts/candidate_design/nanomelt_yield_classification_v2_20260819"
NANOBERT_RESULT_DIR = ROOT / "docs/result_artifacts/candidate_design/nanobert_yield_validation_result_20260813"
TNP_RESULT_DIR = ROOT / "docs/result_artifacts/candidate_design/tnp_yield_validation_result_v2_20260814"
RP3NET_RESULT_DIR = ROOT / "docs/result_artifacts/candidate_design/rp3net_yield_validation_result_20260818"
PLM_SOL_RESULT_DIR = ROOT / "docs/result_artifacts/candidate_design/plm_sol_yield_validation_result_20260819"
FIXED5_RESULT_DIR = ROOT / "docs/result_artifacts/candidate_design/fixed5mg_predictor_classification_20260819"

REPORT = OUT / "Nb252_BL21_expression_optimization_stage_report_2026_W31_W34.docx"
GUIDE = OUT / "Nb252_predictor_and_selection_guide.docx"

BLACK = "000000"
GRAY = "555555"
FONT = "Microsoft YaHei"


def _read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def _write_csv_bom(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore", lineterminator="\n")
        writer.writeheader()
        writer.writerows(rows)


def _set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = FONT
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for slot in ("eastAsia", "ascii", "hAnsi", "cs"):
        r_fonts.set(qn(f"w:{slot}"), FONT)
    for theme_slot in ("eastAsiaTheme", "asciiTheme", "hAnsiTheme", "cstheme"):
        r_fonts.attrib.pop(qn(f"w:{theme_slot}"), None)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    for name, size, bold in [
        ("Normal", 10.5, False),
        ("Title", 21, True),
        ("Subtitle", 11.5, False),
        ("Heading 1", 15, True),
        ("Heading 2", 12.5, True),
        ("Caption", 9, False),
    ]:
        style = doc.styles[name]
        style.font.name = FONT
        r_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        for slot in ("eastAsia", "ascii", "hAnsi", "cs"):
            r_fonts.set(qn(f"w:{slot}"), FONT)
        for theme_slot in ("eastAsiaTheme", "asciiTheme", "hAnsiTheme", "cstheme"):
            r_fonts.attrib.pop(qn(f"w:{theme_slot}"), None)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.size = Pt(size)
        style.font.bold = bold
    normal = doc.styles["Normal"].paragraph_format
    normal.line_spacing = 1.25
    normal.space_after = Pt(5)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(4)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(7)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(3)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(5)


def _set_cell_text(cell, text: object, *, bold: bool = False, size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    _set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_cell_margins(cell, *, vertical: int, horizontal: int = 85) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", vertical), ("bottom", vertical), ("start", horizontal), ("end", horizontal)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_border(container, edge: str, *, value: str, size: int = 0) -> None:
    node = container.find(qn(f"w:{edge}"))
    if node is None:
        node = OxmlElement(f"w:{edge}")
        container.append(node)
    node.set(qn("w:val"), value)
    node.set(qn("w:sz"), str(size))
    node.set(qn("w:space"), "0")
    node.set(qn("w:color"), BLACK)


def _apply_three_line_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    _set_border(borders, "top", value="single", size=12)
    _set_border(borders, "bottom", value="single", size=12)
    for edge in ("left", "right", "insideH", "insideV"):
        _set_border(borders, edge, value="nil")

    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        _set_border(tc_borders, "bottom", value="single", size=8)


def _apply_table_geometry(doc: Document, table, weights: list[float] | None) -> None:
    section = doc.sections[-1]
    total_twips = int((section.page_width - section.left_margin - section.right_margin) / 635)
    if weights is None:
        weights = [1.0] * len(table.columns)
    if len(weights) != len(table.columns) or any(weight <= 0 for weight in weights):
        raise ValueError("Table widths must contain one positive weight per column")
    scale = total_twips / sum(weights)
    widths = [int(round(weight * scale)) for weight in weights]
    widths[-1] += total_twips - sum(widths)

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Twips(width)
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def _add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[object]],
    widths: list[float] | None = None,
    size: float = 8.5,
    vertical_padding: int = 55,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[idx], header, bold=True, size=size)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            _set_cell_text(cells[idx], value, size=size)
    for row in table.rows:
        for cell in row.cells:
            _set_cell_margins(cell, vertical=vertical_padding)
    _apply_table_geometry(doc, table, widths)
    _apply_three_line_borders(table)
    return table


def _iter_document_paragraphs(doc: Document):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def _normalize_document_fonts(doc: Document) -> None:
    for paragraph in _iter_document_paragraphs(doc):
        for run in paragraph.runs:
            _set_run_font(run)


def _add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    _set_run_font(run, size=10.2)


def _add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    _set_run_font(run, size=8.8)


def _add_figure(doc: Document, path: Path, caption: str, width: float = 6.6) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(path), width=Inches(width))
    _add_caption(doc, caption)


def _fmt(value: str, digits: int = 3) -> str:
    return f"{float(value):.{digits}f}"


def _copy_source_figures() -> dict[str, Path]:
    FIG.mkdir(parents=True, exist_ok=True)
    sources = {
        "identity": ROOT / "docs/result_artifacts/input_baseline/summary/input_baseline_qc.png",
        "constraint_tracks": CONSERVATION_DIR / "nb252_conservation_constraint_tracks.png",
        "natural_logo": LOGO_DIR / "nb252_neighbor_sequence_logo_with_regions.png",
        "project_logo": LOGO_DIR / "project_expression_vhh_sequence_logo_with_regions.png",
        "landscape": LANDSCAPE_DIR / "expression_single_mutant_landscape.png",
        "scatter": LANDSCAPE_DIR / "expression_single_mutant_scatter.png",
        "selection": SELECTION_DIR / "expression_single_mutant_selection.png",
    }
    copied: dict[str, Path] = {}
    for key, source in sources.items():
        if not source.is_file():
            raise FileNotFoundError(source)
        target = FIG / f"source_{key}.png"
        shutil.copy2(source, target)
        copied[key] = target
    return copied


def _configure_plot_font() -> None:
    candidates = [Path(r"C:\Windows\Fonts\msyh.ttc"), Path(r"C:\Windows\Fonts\simhei.ttf")]
    font_path = next((p for p in candidates if p.exists()), None)
    if font_path:
        font_manager.fontManager.addfont(str(font_path))
        plt.rcParams["font.family"] = font_manager.FontProperties(fname=str(font_path)).get_name()
    plt.rcParams["axes.unicode_minus"] = False


def _make_route_figure() -> Path:
    _configure_plot_font()
    path = FIG / "figure_current_route.png"
    fig, ax = plt.subplots(figsize=(11.5, 2.6), dpi=220)
    ax.axis("off")
    labels = [
        "权威128-aa亲本\n+ 24界面位点冻结",
        "天然VHH保守性\n+ 80位置不可常规突变",
        "847条约束单突\n+ 单突且保留SSGS",
        "NetSolP / NanoMelt\n+ AntiFold统一评价",
        "幅度分档\n+ 风险与位点多样性",
        "30条计算试选\n+ 等待BL21实验",
    ]
    xs = [0.08, 0.25, 0.42, 0.59, 0.76, 0.93]
    for idx, (x, label) in enumerate(zip(xs, labels)):
        ax.text(x, 0.52, label, ha="center", va="center", fontsize=9.5,
                bbox=dict(boxstyle="round,pad=0.55", facecolor="white", edgecolor="black", linewidth=1.1))
        if idx < len(xs) - 1:
            ax.annotate("", xy=(xs[idx + 1] - 0.075, 0.52), xytext=(x + 0.075, 0.52),
                        arrowprops=dict(arrowstyle="->", color="black", lw=1.2))
    ax.text(0.5, 0.06, "硬约束先于性质评价；当前30条为待复核的单突实验假设。",
            ha="center", va="center", fontsize=9.2)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def _metric_row(path: Path, feature: str) -> dict[str, str]:
    rows = _read_csv(path)
    matches = [row for row in rows if row.get("feature") == feature]
    if len(matches) != 1:
        raise ValueError(f"Expected one row for {feature} in {path}, found {len(matches)}")
    return matches[0]


def _scheme_row(path: Path, scheme: str) -> dict[str, str]:
    rows = _read_csv(path)
    matches = [row for row in rows if row.get("outer_scheme") == scheme]
    if len(matches) != 1:
        raise ValueError(f"Expected one {scheme} row in {path}, found {len(matches)}")
    return matches[0]


def _make_tool_validation_figure() -> Path:
    _configure_plot_font()
    path = FIG / "figure_tool_yield_validation.png"

    continuous_specs = [
        ("NetSolP U", NETSOLP_RESULT_DIR / "netsolp_yield_associations.csv", "predicted_usability"),
        ("NetSolP S", NETSOLP_RESULT_DIR / "netsolp_yield_associations.csv", "predicted_solubility"),
        ("NanoMelt Tm", NANOMELT_RESULT_DIR / "nanomelt_yield_associations.csv", "nanomelt_predicted_apparent_tm_c"),
        ("nanoBERT PLL", NANOBERT_RESULT_DIR / "nanobert_yield_associations.csv", "nanobert_mean_pll_raw"),
        ("TNP PSH", TNP_RESULT_DIR / "tnp_yield_associations.csv", "tnp_psh"),
        ("RP3Net", RP3NET_RESULT_DIR / "rp3net_yield_associations.csv", "rp3net_expression_probability"),
        ("PLM_Sol", PLM_SOL_RESULT_DIR / "plm_sol_yield_associations.csv", "plm_sol_solubility_score"),
    ]
    continuous = []
    for label, source, feature in continuous_specs:
        row = _metric_row(source, feature)
        continuous.append(
            {
                "tool_metric": label,
                "numeric_n": int(row["numeric_n"]),
                "stratified_spearman_rho": float(row["stratified_spearman_rho"]),
            }
        )

    fixed5_rows = _read_csv(FIXED5_RESULT_DIR / "fixed5mg_classification_metrics.csv")
    fixed5 = []
    for label in ["NetSolP U", "NetSolP S", "RP3Net"]:
        matches = [
            row for row in fixed5_rows
            if row["predictor"] == label and row["outer_scheme"] == "leave_one_cluster_out"
        ]
        if len(matches) != 1:
            raise ValueError(f"Expected one fixed-5-mg cluster row for {label}")
        fixed5.append(matches[0])
    plm_fixed5 = _scheme_row(
        PLM_SOL_RESULT_DIR / "plm_sol_fixed5mg_metrics.csv", "leave_one_cluster_out"
    )
    plm_fixed5 = {**plm_fixed5, "predictor": "PLM_Sol"}
    fixed5.append(plm_fixed5)

    nanomelt_class = _scheme_row(
        NANOMELT_CLASS_DIR / "nanomelt_yield_classification.csv", "leave_one_cluster_out"
    )
    summary_rows: list[dict[str, object]] = []
    fixed5_by_label = {row["predictor"]: row for row in fixed5}
    for row in continuous:
        label = str(row["tool_metric"])
        classification = fixed5_by_label.get(label)
        if label == "NanoMelt Tm":
            classification = nanomelt_class
            label_rule = "在每个训练折内按数据来源中位数分组"
        elif classification:
            label_rule = "原始记录产量 >=5 mg/L"
        else:
            label_rule = "未评价"
        summary_rows.append(
            {
                **row,
                "classification_label": label_rule,
                "classification_n": classification.get("n", "") if classification else "",
                "roc_auc": classification.get("roc_auc", "") if classification else "",
                "pr_auc": classification.get("pr_auc_average_precision", "") if classification else "",
                "mcc": classification.get("mcc", "") if classification else "",
                "balanced_accuracy": classification.get("balanced_accuracy", "") if classification else "",
            }
        )
    summary_rows.append(
        {
            "tool_metric": "AntiFold experimental-complex view",
            "numeric_n": 1,
            "stratified_spearman_rho": "",
            "classification_label": "not applicable: only Nb252 has matched experimental complex",
            "classification_n": "",
            "roc_auc": "",
            "pr_auc": "",
            "mcc": "",
            "balanced_accuracy": "",
        }
    )
    _write_csv_bom(
        TABLE / "tool_yield_validation_summary.csv",
        summary_rows,
        [
            "tool_metric", "numeric_n", "stratified_spearman_rho", "classification_label",
            "classification_n", "roc_auc", "pr_auc", "mcc", "balanced_accuracy",
        ],
    )

    fig, axes = plt.subplots(2, 2, figsize=(12.2, 8.0), dpi=220)
    ax = axes[0, 0]
    labels = [str(row["tool_metric"]) for row in continuous]
    rhos = [float(row["stratified_spearman_rho"]) for row in continuous]
    ns = [int(row["numeric_n"]) for row in continuous]
    colors = ["#2F75B5" if label in {"NetSolP U", "NetSolP S"} else "#7F7F7F" for label in labels]
    y = list(range(len(labels)))
    ax.barh(y, rhos, color=colors)
    ax.axvline(0, color="black", linewidth=0.8)
    ax.set_yticks(y, [f"{label} (n={n})" for label, n in zip(labels, ns)])
    ax.invert_yaxis()
    ax.set_xlim(-0.6, 0.6)
    ax.set_xlabel("来源分层 Spearman ρ")
    ax.set_title("A  连续指标与原始记录产量的来源分层关联")
    for yi, value in zip(y, rhos):
        ax.text(value + (0.02 if value >= 0 else -0.02), yi, f"{value:.3f}",
                ha="left" if value >= 0 else "right", va="center", fontsize=8)
    ax.spines[["top", "right"]].set_visible(False)

    fixed_labels = [row["predictor"] for row in fixed5]
    x = list(range(len(fixed_labels)))
    width = 0.34
    ax = axes[0, 1]
    ax.bar([value - width / 2 for value in x], [float(row["roc_auc"]) for row in fixed5],
           width, label="ROC-AUC", color="#2F75B5")
    ax.bar([value + width / 2 for value in x], [float(row["pr_auc_average_precision"]) for row in fixed5],
           width, label="PR-AUC", color="#A5A5A5")
    for xi, row in zip(x, fixed5):
        ax.text(xi - width / 2, float(row["roc_auc"]) + 0.018, f"{float(row['roc_auc']):.3f}",
                ha="center", va="bottom", fontsize=7)
        ax.text(xi + width / 2, float(row["pr_auc_average_precision"]) + 0.018,
                f"{float(row['pr_auc_average_precision']):.3f}", ha="center", va="bottom", fontsize=7)
    ax.axhline(0.5, color="black", linestyle="--", linewidth=0.8)
    ax.set_xticks(x, fixed_labels, rotation=18, ha="right")
    ax.set_ylim(0, 1)
    ax.set_title("B  固定5 mg/L展示：序列簇留一AUC")
    ax.legend(frameon=False, fontsize=8)
    ax.spines[["top", "right"]].set_visible(False)

    ax = axes[1, 0]
    ax.bar([value - width / 2 for value in x], [float(row["mcc"]) for row in fixed5],
           width, label="MCC", color="#2F75B5")
    ax.bar([value + width / 2 for value in x], [float(row["balanced_accuracy"]) for row in fixed5],
           width, label="Balanced accuracy", color="#A5A5A5")
    for xi, row in zip(x, fixed5):
        ax.text(xi - width / 2, float(row["mcc"]) + 0.018, f"{float(row['mcc']):.3f}",
                ha="center", va="bottom", fontsize=7)
        ax.text(xi + width / 2, float(row["balanced_accuracy"]) + 0.018,
                f"{float(row['balanced_accuracy']):.3f}", ha="center", va="bottom", fontsize=7)
    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_xticks(x, fixed_labels, rotation=18, ha="right")
    ax.set_ylim(-0.1, 1)
    ax.set_title("C  固定5 mg/L展示：序列簇留一分类性能")
    ax.legend(frameon=False, fontsize=8)
    ax.spines[["top", "right"]].set_visible(False)

    ax = axes[1, 1]
    ax.axis("off")
    ax.set_title("D  当前保留工具及其证据角色", loc="left")
    boxes = [
        (0.95, "NetSolP", "S/U与产量有有限关联；在序列簇外验证中不弱于PLM_Sol，保留为可溶性软偏好"),
        (0.63, "NanoMelt", f"产量分类AUC={float(nanomelt_class['roc_auc']):.3f}、MCC={float(nanomelt_class['mcc']):.3f}；仅作稳定性约束"),
        (0.31, "AntiFold", "47条中仅Nb252有匹配实验复合物，产量分类不适用；保留为结构相容性约束"),
    ]
    for y0, title, text in boxes:
        ax.text(0.03, y0, title, transform=ax.transAxes, fontsize=10, fontweight="bold", va="top",
                bbox=dict(boxstyle="round,pad=0.45", facecolor="#D9EAF7", edgecolor="#2F75B5"))
        ax.text(0.28, y0 - 0.01, text, transform=ax.transAxes, fontsize=8.6, va="top", wrap=True)
    ax.text(
        0.03, 0.06,
        "PLM_Sol分层相关虽较高，但与NetSolP高度重叠；加入NetSolP S后，序列簇外Spearman反降0.140。\nRP3Net、nanoBERT和TNP也未显示独立样本外增量，故不进入候选排序。",
        transform=ax.transAxes, fontsize=7.8, va="bottom",
    )

    fig.text(
        0.5, 0.012,
        "B/C仅为固定5 mg/L的展示性比较；预测阈值在外层训练折内按MCC选择。NanoMelt使用来源内中位数标签，不能与B/C直接横比；AntiFold分类不适用。",
        ha="center", fontsize=8,
    )
    fig.tight_layout(rect=(0, 0.04, 1, 1))
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def _make_funnel_figure() -> Path:
    _configure_plot_font()
    path = FIG / "figure_selection_funnel.png"
    labels = ["约束单突空间", "幅度短名单", "规则合格", "计算试选"]
    values = [847, 40, 37, 30]
    fig, ax = plt.subplots(figsize=(7.5, 3.5), dpi=220)
    y = list(range(len(labels)))
    ax.barh(y, values, color=["#8C8C8C", "#686868", "#4A4A4A", "#222222"])
    ax.set_yticks(y, labels)
    ax.invert_yaxis()
    ax.set_xlabel("候选数")
    ax.set_title("847条单突到30条计算试选的筛选漏斗")
    for yi, value in zip(y, values):
        ax.text(value + 8, yi, str(value), va="center", fontsize=10)
    ax.spines[["top", "right"]].set_visible(False)
    ax.set_xlim(0, 910)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def _trial_reason(row: dict[str, str]) -> str:
    tier = row["selection_tier"]
    mapping = {
        "A_multi_family": "至少两个预测家族达到中等/强改善，且无中等以上恶化",
        "B_single_family_strong": "一个预测家族强改善，其余指标未出现明确恶化",
        "C_single_family_moderate": "一个预测家族中等改善，其余指标整体可接受",
        "D_controlled_tradeoff": "存在单项中等恶化，但有独立中等改善并保留机制多样性",
        "E_stable_word_exploratory": "稳定词新增假设对照；四项性质仅弱/近中性证据",
    }
    return mapping[tier]


def _short_mutation_label(row: dict[str, str]) -> str:
    return row["mutation_reported_label"].split()[-1]


def _candidate_summary_rows(rows: list[dict[str, str]]) -> list[list[object]]:
    output = []
    for row in rows:
        source = "实验复合物" if row["antifold_selection_source"] == "experimental_complex_context" else "AF3单体补充"
        output.append([
            row["trial_selection_order"], _short_mutation_label(row), row["region"],
            row["selection_tier"].split("_", 1)[0], source, _trial_reason(row),
        ])
    return output


def _candidate_metric_rows(rows: list[dict[str, str]]) -> list[list[object]]:
    stable_word_labels = {"unchanged": "不变", "gain_only": "新增"}
    output = []
    for row in rows:
        output.append([
            _short_mutation_label(row), row["region"],
            _fmt(row["netsolp_delta_usability_vs_current_wt"]),
            _fmt(row["netsolp_delta_solubility_vs_current_wt"]),
            _fmt(row["nanomelt_delta_predicted_apparent_tm_c_vs_current_wt"], 2),
            _fmt(row["antifold_selection_delta_log_probability"]),
            stable_word_labels.get(row["stable_word_effect"], row["stable_word_effect"]),
            row["selection_tier"].split("_", 1)[0],
        ])
    return output


def _add_page_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    _set_run_font(r, size=21, bold=True)
    if subtitle:
        p2 = doc.add_paragraph(style="Subtitle")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        _set_run_font(r2, size=11.5)


def _add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Nb252 BL21表达优化阶段报告｜计算结果仅用于实验优先级排序")
        _set_run_font(run, size=8)


def _build_report(trial: list[dict[str, str]], reserve: list[dict[str, str]], figures: dict[str, Path]) -> None:
    doc = Document()
    _style_document(doc)
    _add_page_title(doc, "Nb252纳米抗体BL21表达优化阶段报告")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("当前版本聚焦表达量优化：冻结实验界面、避开天然保守位点、仅交付单突假设。")
    _set_run_font(r, size=13, bold=True)
    doc.add_paragraph()
    _add_figure(doc, figures["route"], "图1  Nb252表达优化的约束、评价与单突试选流程。", 6.8)

    doc.add_page_break()
    doc.add_heading("执行摘要", level=1)
    _add_bullet(doc, "目标：在不突变实验界面、关键保守位点、二硫键Cys22/Cys95及末端SSGS的前提下，提出30条Nb252单突序列用于BL21表达实验。")
    _add_bullet(doc, "设计空间：天然VHH保守性和结构保护规则联合冻结完整序列中的80个位置；47个常规可扫位置产生846条非Cys单突，另加Q5V共识回变，共847条。")
    _add_bullet(doc, "工具：NetSolP提供S/U，NanoMelt提供预测表观Tm，AntiFold提供结构条件下的逐位点氨基酸兼容性。三者均是预测信号，不等同于实测产量。")
    _add_bullet(doc, "筛选：原始小数保留用于追溯，但候选选择只使用预定义幅度档位。847条中40条进入幅度短名单，排除2条序列风险和2条多重中等恶化后得到36条规则候选；另加入T99F稳定词假设对照，形成37条可选池，最终试选30条、候补7条。")
    _add_bullet(doc, "当前结论：30条仍是计算试选而非最终实验释放。F30（9条）、Q1（5条）和T27（4条）存在明显位点集中，需由项目负责人确认是否接受该探索密度。")
    doc.add_heading("当前试选面板概览", level=2)
    _add_table(doc, ["项目", "结果"], [
        ["候选形式", "30条单氨基酸替换；均为128 aa并保留末端SSGS"],
        ["候选构成", "26条严格核心 + 3条受控权衡 + 1条稳定词探索（T99F）"],
        ["位点覆盖", "13个完整序列位置；CDR 19条，FR 11条"],
        ["AntiFold证据", "24条采用实验复合物视图；6条实验缺失坐标位置采用独立AF3单体补充视图"],
        ["释放状态", "待用户/导师复核，尚未称为最终实验面板"],
    ], widths=[1.45, 5.2], size=9)

    doc.add_page_break()
    doc.add_heading("1. 项目目标与设计原则", level=1)
    doc.add_paragraph("本项目以Nb252在BL21体系中的表达量优化为目标。在保持实验界面、天然保守位点、二硫键和末端SSGS不变的前提下，对允许位置的完整单突空间进行表达相关性质评价，并形成可解释、可实验验证的30条单突试选。")
    doc.add_heading("1.1 当前设计原则", level=2)
    for text in [
        "实验结构中定义的全部VHH界面残基冻结，不因其他预测性质而开放。",
        "不把FR与CDR机械划分为表达区和结合区；只要不属于硬冻结集合，FR/CDR均可被评价。",
        "最终实验阶段先测试单突；任何组合突变均等待单突实验结果。",
        "工具信号以互补证据和风险约束使用，不以简单加权总分或微小小数差异驱动选择。",
    ]:
        _add_bullet(doc, text)

    doc.add_page_break()
    doc.add_heading("2. 权威输入与结构保护边界", level=1)
    doc.add_paragraph("设计亲本为合作者确认的128-aa Nb252完整序列。末端SSGS不是连接肽，而是构建体组成部分；完整序列位置125–128在所有候选中保持不变。实验复合物NK2R–Nb252是结合姿态和界面保护的主要结构证据；AF3 VHH仅补充实验未建模区的预测视图。")
    _add_table(doc, ["输入/证据", "本项目用途", "不能据此声称"], [
        ["Nb252 128-aa亲本", "定义WT、候选序列与末端SSGS", "不能自行裁剪或改写构建体"],
        ["实验NK2R–Nb252复合物", "界面保护、实验坐标可评价范围", "界面残基不是能量热点排名"],
        ["AF3 VHH预测", "实验缺失区的独立结构兼容性参考", "不能替代实验CDR3或受体接触"],
        ["47条BL21产量序列", "验证预测指标是否与产量有关", "不能训练高容量、可外推的表达模型"],
    ], widths=[1.55, 2.5, 2.55], size=8.7)
    doc.add_heading("2.1 不可突变集合", level=2)
    _add_bullet(doc, "24个实验界面残基：由严格<4.0 Å聚合物重原子中心距离定义复现。")
    _add_bullet(doc, "54个Nb252自身等于天然优势残基的高置信保守位点。")
    _add_bullet(doc, "Cys22/Cys95、末端SSGS以及其他合同规定的不可变位置。")
    _add_bullet(doc, "Q5不称为亲本保守位点：天然共识为V，因此只开放Q5V共识回变。")
    _add_figure(doc, figures["identity"], "图2  序列、实验坐标覆盖、AF3坐标覆盖及实验界面位置基线。灰色代表实验坐标不可评价，不等同于序列缺失。", 6.8)

    doc.add_page_break()
    doc.add_heading("3. 天然VHH保守性与可变位置", level=1)
    doc.add_paragraph("保守性参考采用TNP研究公开的4,059条非冗余天然VHH序列。ANARCII按照IMGT抗体标准编号体系完成位置对应与完整性审核，4,057条通过；先把完整IMGT域相似度达到90%的近重复序列归为同一簇，得到3,784个有效簇，再依据框架区序列相似度和覆盖率定义Nb252邻域。邻域包含1,564条序列、1,532个有效簇。按簇等权可避免大量近重复序列放大某一克隆家族。FR表示相对保守的框架区，CDR表示互补决定区。")
    doc.add_heading("3.1 高置信保守位点的定义", level=2)
    _add_bullet(doc, "邻域优势残基频率≥0.90、位置覆盖率≥0.80、有效簇数≥50。")
    _add_bullet(doc, "全局与邻域优势残基一致，且全局优势频率≥0.80。")
    _add_bullet(doc, "只有当Nb252自身残基等于该共同优势残基时，才因天然保守性完全冻结。")
    _add_bullet(doc, "满足统计保守性但Nb252偏离共识的位置不任意开放；仅允许共识回变。")
    _add_figure(doc, figures["constraint_tracks"], "图3  Nb252天然保守性与设计约束轨道。硬保守、界面、二硫键和末端构成设计禁区。", 6.8)

    doc.add_page_break()
    doc.add_heading("3.2 序列Logo提供的两类参照", level=1)
    doc.add_paragraph("天然邻域Logo反映经过簇权重校正的Nb252邻近VHH序列规律；项目47条表达序列Logo反映现有样本集合内部的残基分布。后者不能替代天然保守性合同，因为样本来源、构建体范围和序列完整性并不完全一致，且样本量较小。")
    _add_figure(doc, figures["natural_logo"], "图4  Nb252天然邻域序列Logo，按IMGT FR/CDR区域标注。", 6.7)
    _add_figure(doc, figures["project_logo"], "图5  项目47条表达序列中的可编号H链Logo，仅用于描述项目样本序列空间。", 6.7)

    doc.add_page_break()
    doc.add_heading("4. 847条约束单突空间", level=1)
    doc.add_paragraph("约束合并后，完整128-aa序列中有80个位置不可常规突变。其余47个位置对除Cys外的19种替换进行完整枚举，共846条；加上Q5V共识回变，总计847条。每条候选均为唯一单突、长度128 aa、保留末端SSGS、不引入新Cys。")
    _add_table(doc, ["步骤", "数量", "说明"], [
        ["完整序列位置", 128, "权威亲本序列索引"],
        ["联合冻结位置", 80, "界面、天然保守、二硫键、末端及合同保护"],
        ["常规可扫描位置", 47, "每个位置19种非Cys替换"],
        ["常规单突", 846, "47×19，全部为单突"],
        ["共识回变", 1, "Q5V，仅此一条"],
        ["完整评价空间", 847, "进入统一性质计算，不是最终候选数"],
    ], widths=[1.7, 1.0, 3.9], size=9)
    doc.add_heading("4.1 为什么不先按FR/CDR切割", level=2)
    doc.add_paragraph("表达量可能同时受表面电荷、疏水性、局部构象稳定和折叠兼容性影响，这些因素跨越FR/CDR边界。因此本轮只使用结构界面和保守性作为硬约束；其余可变位置统一进入同一性质评价框架。")
    _add_figure(doc, figures["funnel"], "图6  完整约束空间到计算试选的数量漏斗。", 6.3)

    doc.add_page_break()
    doc.add_heading("5. 工具选择与证据角色", level=1)
    doc.add_paragraph("47条序列的BL21产量数据用于判断工具指标是否具备项目内解释价值。验证同时查看连续关联和预设方向下的样本外分类表现；没有任何工具被证明可直接预测未测Nb252突变体的mg/L产量。")
    _add_table(doc, ["工具/指标", "含义与方向", "当前证据角色"], [
        ["NetSolP S", "预测可溶性得分；越高通常越有利", "有限产量关联；作为可溶性软偏好"],
        ["NetSolP U", "模型定义的综合可用性得分；0–1，越高通常越有利", "与S同属一个模型家族，不重复计票"],
        ["NanoMelt Tm", "预测表观熔解温度（°C）；越高通常代表更稳定", "作为稳定性约束，不称为产量排序器"],
        ["AntiFold ΔlogP", "结构条件下突变残基相对WT的对数概率变化；越高越兼容", "实验复合物视图优先；实验缺失位置用独立AF3单体补充"],
        ["稳定词", "简并字母表下新增连续片段", "未验证为产量指标，仅保留一个探索对照T99F"],
    ], widths=[1.25, 2.55, 2.8], size=8.6)
    doc.add_heading("5.1 多工具与产量数据的验证", level=2)
    doc.add_paragraph("连续分析优先使用来源分层Spearman秩相关：先在各数据来源内按相对次序计算关联，再合并来源证据，以降低LTT和WCC两组数据来源差异的混杂。Spearman ρ范围为−1至1，正值表示指标随产量总体上升。固定5 mg/L的分类图只纳入31条具有个体数值产量的序列；把相似度达到90%的近重复序列归为同一簇，每次完整留出一个序列簇，可减少近同源序列同时出现在训练和测试中的信息泄漏。每个外层训练折内按最大MCC选择预测阈值，再在留出的序列簇上评价。NanoMelt因有效数值样本为27条且采用来源内中位数标签，单独报告，不能与固定5 mg/L面板直接横向比较。")
    _add_figure(doc, figures["tool_validation"], "图7  已测试工具指标与BL21原始记录产量的连续关联、展示性分类性能及当前证据角色。A为来源分层Spearman秩相关；B/C为固定5 mg/L、序列簇留一结果；D说明保留NetSolP、NanoMelt和AntiFold的不同理由。", 6.9)
    doc.add_paragraph("分类指标中，ROC-AUC和PR-AUC分别概括不同阈值下的受试者工作特征和精确率-召回率表现，越接近1越好；MCC综合真阳性、真阴性、假阳性和假阴性，范围为−1至1，0附近相当于无稳定判别；平衡准确率是灵敏度与特异度的平均值。图中的nanoBERT PLL表示抗体语言模型对序列自然性的对数似然，TNP PSH表示预测表面疏水性，RP3Net表示大肠杆菌表达概率，PLM_Sol表示预测可溶性得分。")
    doc.add_paragraph("NetSolP U/S与产量呈有限正关联，且固定5 mg/L展示中S的序列簇留一ROC-AUC为0.777、MCC为0.411；因此NetSolP保留为可溶性软偏好，而不是mg/L预测器。NanoMelt的来源分层Spearman为0.040，来源内中位数分类ROC-AUC为0.571、MCC为0.408，连续证据不支持产量排序，但其Tm指标仍提供与可溶性不同的稳定性约束。AntiFold无法在47条数据上进行公平的产量相关或分类验证，因为只有Nb252具有匹配实验复合物；它保留的依据是对目标复合物逐位点突变的结构相容性约束。")
    doc.add_paragraph("PLM_Sol的来源分层Spearman为0.473，确实高于NetSolP U的0.394和NetSolP S的0.376；但这只是同一批31条数值记录上的单变量相关，不能单独决定工具取舍。两种可溶性工具的输出高度重叠：PLM_Sol与NetSolP U/S在47条序列上的Spearman分别为0.793和0.640，在31条数值记录中进一步升至0.926和0.715。因此把两者同时用于筛选，主要是在重复计算同类证据，而不是增加独立信息。")
    doc.add_paragraph("更关键的是样本外比较没有显示PLM_Sol优于NetSolP。固定5 mg/L的序列簇留一结果中，PLM_Sol的ROC-AUC为0.761，低于NetSolP S的0.777；二者MCC均为0.411、平衡准确率均为0.697。连续产量预测中，NetSolP S单独的序列簇外Spearman为0.455，加入PLM_Sol后反而降至0.315，独立增量为−0.140。也就是说，PLM_Sol较高的来源分层相关没有转化为更好的近同源外泛化，也没有给NetSolP增加有效补充。项目因此保留部署更直接、已覆盖完整单突空间的NetSolP作为可溶性代表，而不让两个高度相关的可溶性模型重复投票；这不表示PLM_Sol本身无效，只表示当前47条数据不支持其替代或叠加NetSolP。")
    doc.add_paragraph("RP3Net、nanoBERT和TNP虽已测试，也没有显示可独立外推到Nb252单突空间的稳定样本外增量，因此不进入本轮847条候选排序。")
    doc.add_heading("5.2 为什么保留三种工具", level=2)
    _add_bullet(doc, "NetSolP：在现有数据中提供最直接但有限的可溶性/产量相关信号；S与U视为同一工具家族。")
    _add_bullet(doc, "NanoMelt：不作为产量预测器，而用于阻止预测稳定性明显下降的候选；其证据角色与NetSolP不同。")
    _add_bullet(doc, "AntiFold：不参与产量分类，而利用目标实验复合物判断突变残基是否与当前VHH结构环境相容。")
    doc.add_heading("5.3 幅度分档而非小数排序", level=2)
    _add_bullet(doc, "保留所有原始S/U/Tm/AntiFold数值，但先映射为强有利、中等有利、弱/近中性、中等不利、强不利。")
    _add_bullet(doc, "同一档内的小数差异不用于决定先后，避免模型噪声和微小波动主导候选。")
    _add_bullet(doc, "NetSolP S与U计作一个预测家族；AntiFold与NanoMelt分别代表结构兼容性与稳定性。")

    doc.add_page_break()
    doc.add_heading("6. 847条候选的统一性质评价", level=1)
    doc.add_paragraph("NetSolP和NanoMelt覆盖全部847条。AntiFold中，721条候选所在位置在实验VHH复合物坐标中可评价，优先使用实验复合物视图；126条位于实验结构缺失区，在实验视图中明确标记为“实验坐标不可评价”，并单独使用AF3单体预测结果作为补充。两种来源在图表和候选清单中明确区分。")
    _add_figure(doc, figures["landscape"], "图8  847条约束单突的四指标热图。AntiFold面板中126条实验坐标不可评价项以AF3单体预测结果补充，并保留来源标记。", 6.9)
    doc.add_paragraph("热图显示，大多数单突的NetSolP S/U和NanoMelt Tm变化很小。当前流程因此不把极小变化解释为真实改善；只有达到预定义中等或强幅度时才形成主要入选证据。")

    doc.add_page_break()
    doc.add_heading("6.1 指标之间的关系", level=1)
    doc.add_paragraph("NetSolP S/U、NanoMelt Tm和AntiFold ΔlogP描述不同方面，并不要求同步变化。散点图用于识别明显权衡、孤立极端值和实验坐标缺失位置，而不是拟合一条统一最优线。")
    _add_figure(doc, figures["scatter"], "图9  四指标关键组合散点图。点形区分AntiFold实验复合物视图与AF3单体补充视图；星号标记稳定词新增。", 6.8)
    doc.add_heading("6.2 稳定词作为探索性软偏好", level=2)
    doc.add_paragraph("1,336条稳定词按固定12符号简并字母表进行大小写敏感、允许重叠的精确连续子串匹配。847条单突中22条新增24个命中，825条不变，没有减少项。但在47条产量数据中，稳定词特征未显示稳定正相关或可靠分类能力，因此不作为硬筛选器，只保留T99F用于实验检验该假设。")

    doc.add_page_break()
    doc.add_heading("7. 从847条到30条：幅度分档筛选", level=1)
    doc.add_paragraph("筛选先看幅度和风险，再看位点与机制多样性。847条中40条满足“无明显恶化且至少一个预测家族达到中等改善”。其中2条引入明确序列风险，2条同时出现两个中等不利指标，均不进入可释放池。剩余36条由26条严格核心和10条受控权衡组成；另将不在40条短名单中的T99F作为稳定词探索例外加入，形成37条可供试选，最终选择30条、保留7条候补。")
    _add_table(doc, ["阶段", "数量", "判定"], [
        ["完整约束空间", 847, "全部满足序列与硬约束"],
        ["幅度短名单", 40, "至少一个中等/强改善，且无明显恶化"],
        ["规则合格", 36, "排除2条序列风险和2条多重中等不利"],
        ["探索例外", 1, "T99F新增稳定词；仅弱/近中性性质证据"],
        ["计算试选", 30, "26严格核心 + 3受控权衡 + T99F"],
        ["候补", 7, "受控权衡备选"],
    ], widths=[1.4, 0.8, 4.5], size=9)
    _add_figure(doc, figures["selection"], "图10  当前30条计算试选：漏斗、位点分布和四指标幅度档位。空心三角表示AntiFold采用AF3单体补充视图；金色星号表示T99F稳定词探索候选。", 6.9)

    doc.add_page_break()
    doc.add_heading("8. 当前30条计算试选", level=1)
    doc.add_paragraph("30条覆盖13个位置，分为A–E五类。A/B/C为严格核心，D为受控权衡，E为稳定词探索。表格从上到下体现分层优先级：A优先于B，B优先于C，随后是D和E；但同一层内部不是1–30的精确效力排名，同档小数差异没有用于排序。完整序列与原始预测值见交付FASTA、CSV和附录。")
    for start in range(0, 30, 10):
        if start:
            doc.add_page_break()
            doc.add_heading(f"8. 当前30条计算试选（续：{start + 1}–{min(start + 10, 30)}）", level=1)
        _add_table(doc, ["展示顺序", "单突", "区段", "层", "AntiFold来源", "主要入选理由"],
                   _candidate_summary_rows(trial[start:start + 10]),
                   widths=[0.55, 0.7, 0.55, 0.35, 1.05, 3.1], size=9.2, vertical_padding=105)

    doc.add_page_break()
    doc.add_heading("9. 位点集中、风险与不确定性", level=1)
    _add_table(doc, ["观察", "当前解释", "需要的决策/实验"], [
        ["F30占9条", "该位置在NetSolP/AntiFold中出现较多可接受替换，规则筛选自然集中", "确认是否接受同一位点的广泛替换探索；若实验容量有限，可保留化学性质差异最大的子集"],
        ["Q1占5条", "N端替换主要由单一预测家族支持", "关注起始端加工及表达构建体上下文，实验中保持同一载体和标签"],
        ["T27占4条、F29占2条", "实验结构缺失，AntiFold使用AF3单体补充视图", "解释时降低结构证据等级；不可称为实验复合物支持"],
        ["T99F稳定词探索", "新增长度5稳定词，但四项性质无中等改善", "作为机制对照，不与A–C层等价"],
        ["工具与产量相关有限", "预测可减少明显风险，但不能保证mg/L提高", "WT同批、多重复、保留全部阴性/中性结果"],
    ], widths=[1.25, 2.55, 2.8], size=8.5)
    doc.add_heading("9.1 风险控制", level=2)
    _add_bullet(doc, "所有30条均未触碰界面、硬保守位置、Cys22/Cys95和末端SSGS，也未引入新Cys。")
    _add_bullet(doc, "筛选检查新Pro、局部疏水/电荷变化、N-糖基化、脱酰胺/异构化和氧化易感残基等序列风险。")
    _add_bullet(doc, "受控权衡候选最多允许单项中等不利；同时出现两个中等不利的候选不进入面板。")
    _add_bullet(doc, "AntiFold的AF3单体视图只补充结构兼容性参考，不把实验缺失区转化为实验可评价区。")

    doc.add_page_break()
    doc.add_heading("10. 下一步实验与决策计划", level=1)
    doc.add_heading("10.1 面板释放前复核", level=2)
    _add_bullet(doc, "确认30条中F30/Q1/T27的位点集中度是否符合本轮探索目的。")
    _add_bullet(doc, "确认3条受控权衡和T99F探索对照是否占用实验名额；必要时从7条候补中一对一替换。")
    _add_bullet(doc, "冻结最终FASTA版本后，不再因同档内小数差异调整顺序。")
    doc.add_heading("10.2 BL21表达实验", level=2)
    _add_bullet(doc, "WT与30条单突使用相同载体、信号肽、标签、诱导、培养、裂解、纯化和定量流程。")
    _add_bullet(doc, "建议至少3个独立生物重复；主终点为纯化后mg/L产量，并记录总表达、可溶上清、沉淀与纯化回收。")
    _add_bullet(doc, "预先定义成功标准，例如相对WT提高幅度、重复一致性和不低于WT的可溶比例。")
    _add_bullet(doc, "保留全部阴性、中性和失败结果，用于下一轮局部规则更新。")
    doc.add_heading("10.3 组合突变的阶段门", level=2)
    doc.add_paragraph("本轮不生成或提交组合突变。只有当两个或多个单突在相同实验条件下分别显示可重复改善，且无明显共同风险时，才进入组合设计；组合后需要重新评价NetSolP、NanoMelt和AntiFold，并重新检查界面/保守/末端合同。")
    _add_table(doc, ["阶段门", "当前状态"], [
        ["结构与界面身份", "通过"],
        ["天然VHH保守性合同", "通过"],
        ["847条单突性质矩阵", "通过"],
        ["30条计算试选", "已生成，待人工复核"],
        ["最终实验面板释放", "未通过"],
        ["组合突变设计", "阻断，等待单突实验"],
    ], widths=[2.6, 4.0], size=9)

    doc.add_page_break()
    doc.add_heading("附录A 30条试选的原始变化值", level=1)
    doc.add_paragraph("下表保留原始模型变化值用于追溯。候选选择采用幅度档位，表内小数不用于同档内精细排序。AntiFold ΔlogP为突变残基相对WT的对数概率变化；正值为预测兼容性提高。")
    for start in range(0, 30, 10):
        if start:
            doc.add_page_break()
            doc.add_heading(f"附录A（续：{start + 1}–{min(start + 10, 30)}）", level=1)
        _add_table(doc, ["单突", "区段", "ΔU", "ΔS", "ΔTm °C", "AntiFold ΔlogP", "稳定词", "层"],
                   _candidate_metric_rows(trial[start:start + 10]),
                   widths=[0.65, 0.55, 0.55, 0.55, 0.65, 1.0, 0.8, 0.35], size=7.7)

    doc.add_page_break()
    doc.add_heading("附录B 7条候补", level=1)
    _add_table(doc, ["单突", "区段", "ΔU", "ΔS", "ΔTm °C", "AntiFold ΔlogP", "候补理由"], [
        [_short_mutation_label(row), row["region"], _fmt(row["netsolp_delta_usability_vs_current_wt"]),
         _fmt(row["netsolp_delta_solubility_vs_current_wt"]), _fmt(row["nanomelt_delta_predicted_apparent_tm_c_vs_current_wt"], 2),
         _fmt(row["antifold_selection_delta_log_probability"]), "受控权衡；用于替换位点集中或探索名额"]
        for row in reserve
    ], widths=[0.65, 0.55, 0.55, 0.55, 0.65, 1.0, 2.2], size=7.8)
    doc.add_heading("附录C 术语边界", level=1)
    _add_bullet(doc, "计算试选：通过当前计算规则、推荐进入人工复核的实验假设。")
    _add_bullet(doc, "最终实验面板：经项目负责人确认并冻结用于实际构建/表达的序列集合。当前尚未达到该状态。")
    _add_bullet(doc, "预测改善：模型输出方向有利，不等同于实测产量、溶解度或Tm提高。")
    _add_bullet(doc, "AF3单体补充视图：实验坐标缺失位置的独立预测结构参考，不是实验复合物证据。")
    _add_bullet(doc, "稳定词探索：用户指定的可解释软偏好，尚未被47条产量数据验证。")

    _add_footer(doc)
    doc.core_properties.title = "Nb252纳米抗体BL21表达优化阶段报告"
    doc.core_properties.subject = "Nb252 BL21表达优化阶段结果"
    doc.core_properties.author = "Antibody_optimization project"
    doc.core_properties.keywords = "Nb252, BL21, expression, NetSolP, NanoMelt, AntiFold"
    _normalize_document_fonts(doc)
    doc.save(REPORT)


def _build_guide() -> None:
    doc = Document()
    _style_document(doc)
    section = doc.sections[0]
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.2)
    _add_page_title(doc, "Nb252候选指标与筛选说明", "一页速览｜用于阅读30条试选清单")
    _add_table(doc, ["指标", "方向", "在本项目中的正确解释"], [
        ["NetSolP U", "越高通常越有利", "0–1的综合可用性模型输出；与S同属一个工具家族，不双重计票"],
        ["NetSolP S", "越高通常越有利", "预测可溶性得分；PLM_Sol分层相关虽较高，但簇外验证不优且无独立增量，故保留NetSolP作为同类代表"],
        ["NanoMelt Tm", "越高通常越稳定", "预测表观Tm（°C）；作为稳定性约束，不是mg/L预测"],
        ["AntiFold ΔlogP", "正值更兼容", "突变残基相对WT的结构条件对数概率变化；实验复合物优先，缺失区用AF3单体补充视图"],
        ["稳定词", "新增仅作探索", "简并片段新增；未验证为产量指标，仅T99F作为对照"],
    ], widths=[1.1, 1.1, 4.45], size=8.2)
    doc.add_heading("筛选逻辑", level=2)
    _add_bullet(doc, "先硬排除：24界面位置、天然硬保守位置、Cys22/Cys95、末端SSGS及其他不可变位置。")
    _add_bullet(doc, "再按幅度档位：强/中等有利、弱/近中性、中等/强不利；同档内小数不排序。")
    _add_bullet(doc, "A–C层为严格核心；D层为单项中等不利的受控权衡；E层T99F是稳定词假设对照。")
    _add_bullet(doc, "表格顺序表示A→B→C→D→E的分层优先级；同层内部不是精确效力排名。")
    _add_bullet(doc, "空心三角：AntiFold使用AF3单体补充视图；金色星号：稳定词探索候选。")
    _add_bullet(doc, "当前30条是计算试选，需人工确认位点集中后才能冻结为实验面板。")
    doc.add_heading("实验判读", level=2)
    p = doc.add_paragraph("所有预测均不是实验结果。建议WT与30条单突同批、至少3个生物重复，以纯化后mg/L为主终点，同时记录总表达、可溶比例和纯化回收；单突结果返回前不组合。")
    for run in p.runs:
        _set_run_font(run, size=9.2)
    _normalize_document_fonts(doc)
    doc.save(GUIDE)


def _write_delivery(trial: list[dict[str, str]], reserve: list[dict[str, str]]) -> None:
    DELIVERY.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SELECTION_DIR / "expression_single_mutant_trial30.fasta", DELIVERY / "Nb252_trial30_single_mutants.fasta")
    shutil.copy2(SELECTION_DIR / "expression_single_mutant_reserve.fasta", DELIVERY / "Nb252_reserve7_single_mutants.fasta")
    fields = [
        "trial_selection_order", "mutation_reported_label", "reported_sequence_index_1based", "region",
        "imgt_position_label", "selection_tier", "trial_selection_reason", "sequence",
        "netsolp_predicted_usability", "netsolp_delta_usability_vs_current_wt",
        "netsolp_predicted_solubility", "netsolp_delta_solubility_vs_current_wt",
        "nanomelt_predicted_apparent_tm_c", "nanomelt_delta_predicted_apparent_tm_c_vs_current_wt",
        "antifold_selection_source", "antifold_selection_delta_log_probability", "stable_word_effect",
        "final_experimental_panel_released",
    ]
    _write_csv_bom(DELIVERY / "Nb252_trial30_single_mutants.csv", trial, fields)
    _write_csv_bom(DELIVERY / "Nb252_reserve7_single_mutants.csv", reserve, fields)
    readme = (
        "# Nb252 BL21表达优化阶段交付包\n\n"
        "本目录面向导师和合作者，包含Nb252表达优化阶段报告、指标说明、30条计算试选和7条候补。\n\n"
        "- `Nb252_BL21_expression_optimization_stage_report_2026_W31_W34.docx/pdf`：完整阶段报告。\n"
        "- `Nb252_predictor_and_selection_guide.docx/pdf`：一页指标与筛选说明。\n"
        "- `Nb252_trial30_single_mutants.fasta/csv`：30条计算试选；尚未等同于最终实验释放。\n"
        "- `Nb252_reserve7_single_mutants.fasta/csv`：7条候补。\n\n"
        "`trial_selection_order`表示分层展示顺序，不是同一层内部的精确效力排名。\n\n"
        "所有候选均为128-aa单突并保留末端SSGS。预测值不是实测表达量、溶解度或Tm。\n"
    )
    (DELIVERY / "README.md").write_text(readme, encoding="utf-8")


def main() -> int:
    if OUT.exists():
        raise FileExistsError(f"Output directory already exists: {OUT}")
    FIG.mkdir(parents=True)
    TABLE.mkdir(parents=True)
    trial = _read_csv(SELECTION_DIR / "expression_single_mutant_trial30.csv")
    reserve = _read_csv(SELECTION_DIR / "expression_single_mutant_reserve.csv")
    if len(trial) != 30 or len(reserve) != 7:
        raise ValueError(f"Unexpected selection counts: trial={len(trial)}, reserve={len(reserve)}")
    if any(row["final_experimental_panel_released"].lower() != "false" for row in trial):
        raise ValueError("Trial rows must remain unreleased")

    figures = _copy_source_figures()
    figures["route"] = _make_route_figure()
    figures["funnel"] = _make_funnel_figure()
    figures["tool_validation"] = _make_tool_validation_figure()
    _build_report(trial, reserve, figures)
    _build_guide()
    _write_delivery(trial, reserve)

    position_counts = Counter(row["reported_sequence_index_1based"] for row in trial)
    manifest = {
        "schema_version": 1,
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "reporting_period": "2026-W31--W34",
        "active_route": "BL21_expression_only_single_mutants",
        "trial_count": 30,
        "reserve_count": 7,
        "unique_trial_positions": len(position_counts),
        "trial_position_counts": dict(sorted(position_counts.items(), key=lambda item: int(item[0]))),
        "final_experimental_panel_released": False,
        "report_docx": REPORT.name,
        "guide_docx": GUIDE.name,
        "delivery_directory": "delivery",
        "source_selection_release": "expression_single_mutant_trial_selection_v2_20260820",
        "excluded_current_route_methods": ["Rosetta", "PyRosetta", "Flex ddG", "affinity ranking", "double mutants"],
    }
    (OUT / "report_manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    shutil.copy2(REPORT, DELIVERY / REPORT.name)
    shutil.copy2(GUIDE, DELIVERY / GUIDE.name)
    print(REPORT)
    print(GUIDE)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
