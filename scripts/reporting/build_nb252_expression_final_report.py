from __future__ import annotations

import argparse
import csv
import json
import shutil
from collections import Counter
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib import font_manager
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[2]
FINAL_DIR = ROOT / "docs/result_artifacts/candidate_design/expression_final_19plus11_panel_20260822"
PARENT_DIR = ROOT / "docs/result_artifacts/candidate_design/expression_single_mutant_parent19_20260822"
TRIAL_DIR = ROOT / "docs/result_artifacts/candidate_design/expression_single_mutant_trial_selection_v2_20260820"
CURRENT_REPORT_DIR = ROOT / "docs/result_artifacts/weekly_report_result/report_2026_W34_nb252_expression_route"

FONT = "Microsoft YaHei"
BLACK = "000000"
GRAY = "555555"
BLUE = "#2F6F9F"

BAND_LABELS = {
    "strong_favorable": "明显有利",
    "moderate_favorable": "中等有利",
    "weak_favorable": "轻度有利",
    "negligible": "近似中性",
    "weak_adverse": "轻度不利",
    "moderate_adverse": "中等不利",
    "strong_adverse": "明显不利",
}

SINGLE_REASON = {
    "F30S": "NetSolP S明显改善，U和AntiFold中等支持",
    "F30A": "NetSolP U/S与AntiFold均为中等支持",
    "F30R": "U和AntiFold中等支持，并保留带正电替换假设",
    "Q1M": "AntiFold明显支持，保留N端疏水替换假设",
    "Q1H": "NetSolP S中等支持，Tm轻度有利",
    "Q1D": "NetSolP S中等支持，保留酸性替换假设",
    "T27F": "AntiFold明显支持，显式接受S轻度不利",
    "T27D": "AntiFold中等支持，无中等或明显性质恶化",
    "T27S": "AntiFold中等支持，无中等或明显性质恶化",
    "Q5V": "天然优势残基共识回变，Tm明显、AntiFold中等支持",
    "A23S": "F30/Q1/T27外位置各留一条；U与AntiFold提供支持",
    "F29T": "保留F29替换，并避开F29S的Tm轻度不利",
    "Y32L": "F30/Q1/T27外位置各留一条；Tm中等支持",
    "E44G": "AntiFold明显支持，并避开E44A的U轻度不利",
    "S55G": "Tm明显、AntiFold中等支持",
    "K86R": "保守碱性替换，AntiFold中等支持",
    "P87T": "AntiFold中等支持；Tm中等不利作为显式权衡",
    "V97S": "AntiFold中等支持；Tm仅轻度不利",
    "T99F": "新增稳定词的探索性单突，不作为多工具改善候选",
}


def _csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def _json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _write_utf8_lf(path: Path, text: str) -> None:
    with path.open("w", encoding="utf-8", newline="\n") as handle:
        handle.write(text)


def _font(run, size: float | None = None, bold: bool | None = None, color: str = BLACK) -> None:
    run.font.name = FONT
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), FONT)
    run.font.color.rgb = RGBColor.from_string(color)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(0.9)
    specs = {
        "Normal": (10.5, False, 0, 5, 1.25),
        "Title": (21, True, 0, 6, 1.0),
        "Subtitle": (12, False, 0, 5, 1.0),
        "Heading 1": (15, True, 8, 7, 1.05),
        "Heading 2": (12.5, True, 6, 5, 1.05),
        "Caption": (9, False, 3, 6, 1.0),
        "List Bullet": (10.5, False, 0, 4, 1.2),
    }
    for name, (size, bold, before, after, spacing) in specs.items():
        style = doc.styles[name]
        style.font.name = FONT
        fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{key}"), FONT)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.size = Pt(size)
        style.font.bold = bold
        fmt = style.paragraph_format
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.line_spacing = spacing
    doc.styles["List Bullet"].paragraph_format.left_indent = Cm(0.65)
    doc.styles["List Bullet"].paragraph_format.first_line_indent = Cm(-0.32)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _border(parent, edge: str, value: str, size: int = 0) -> None:
    node = parent.find(qn(f"w:{edge}"))
    if node is None:
        node = OxmlElement(f"w:{edge}")
        parent.append(node)
    node.set(qn("w:val"), value)
    node.set(qn("w:sz"), str(size))
    node.set(qn("w:color"), BLACK)
    node.set(qn("w:space"), "0")


def _table(doc: Document, headers: list[str], rows: list[list[object]], weights: list[float], size: float = 8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _repeat_header(table.rows[0])
    for idx, value in enumerate(headers):
        _cell(table.rows[0].cells[idx], value, True, size)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            _cell(cells[idx], value, False, size)
    total = 9500
    widths = [round(total * value / sum(weights)) for value in weights]
    widths[-1] += total - sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total)); tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120"); tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Twips(width)
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(width)); tc_w.set(qn("w:type"), "dxa")
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    _border(borders, "top", "single", 12); _border(borders, "bottom", "single", 12)
    for edge in ("left", "right", "insideH", "insideV"):
        _border(borders, edge, "nil")
    for cell in table.rows[0].cells:
        tc_borders = cell._tc.get_or_add_tcPr().find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            cell._tc.get_or_add_tcPr().append(tc_borders)
        _border(tc_borders, "bottom", "single", 8)
    return table


def _cell(cell, value: object, bold: bool, size: float) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    _font(p.add_run(str(value)), size=size, bold=bold)
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, amount in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(amount)); node.set(qn("w:type"), "dxa")


def _point(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    _font(p.add_run(text), size=10.5)


def _figure(doc: Document, path: Path, caption: str, width_cm: float = 16.4) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(cap.add_run(caption), size=9)


def _new_page_heading(doc: Document, text: str) -> None:
    heading = doc.add_heading(text, level=1)
    heading.paragraph_format.page_break_before = True


def _fmt(value: str, digits: int = 3) -> str:
    return f"{float(value):+.{digits}f}"


def _mutation(row: dict[str, str]) -> str:
    label = row.get("mutation_reported_label", "")
    return label.rsplit(" ", 1)[-1] if label else row["mutation_set"]


def _panel_reason(row: dict[str, str]) -> tuple[str, str]:
    if row["candidate_kind"] == "single_mutant":
        mutation = row["mutation_set"]
        return "父单突", SINGLE_REASON[mutation]
    layer = row["evidence_layer"][0]
    family_count = row["favorable_family_count"]
    strong_count = row["strong_favorable_family_count"]
    detail = f"{family_count}个预测器家族达到中等或明显有利"
    if int(strong_count):
        detail += f"，其中{strong_count}个家族明显有利"
    detail += "；无中等或明显性质恶化，并通过多样性约束"
    return f"{layer}层", detail


def _double_layer_counts(audit: list[dict[str, str]]) -> dict[str, dict[str, int]]:
    layers = ["A_three_families", "B_two_families", "C_one_family", "ineligible"]
    total = Counter(row["evidence_layer"] for row in audit)
    eligible = Counter(
        row["evidence_layer"] for row in audit if row["double_selection_eligibility"] == "eligible"
    )
    selected = Counter(
        row["evidence_layer"] for row in audit if row["double_selection_status"] == "selected_final11"
    )
    return {
        "total": {layer: total[layer] for layer in layers},
        "eligible": {layer: eligible[layer] for layer in layers},
        "selected": {layer: selected[layer] for layer in layers},
    }


def _make_figures(
    out: Path,
    parents: list[dict[str, str]],
    doubles: list[dict[str, str]],
    gate: dict,
    audit: list[dict[str, str]],
    single_gate: dict,
) -> dict[str, Path]:
    plt.rcParams.update({"font.family": [FONT, "Arial", "DejaVu Sans"], "axes.unicode_minus": False})
    fig_dir = out / "figures"
    fig_dir.mkdir(parents=True)
    sources = {
        "constraints": CURRENT_REPORT_DIR / "figures/source_constraints.png",
        "validation": CURRENT_REPORT_DIR / "figures/source_validation.png",
        "landscape": CURRENT_REPORT_DIR / "figures/source_landscape.png",
        "final": FINAL_DIR / "expression_final_panel_overview.png",
    }
    for key, source in sources.items():
        target = fig_dir / f"source_{key}.png"
        shutil.copy2(source, target)
        sources[key] = target

    counts = Counter(row["reported_sequence_index_1based"] for row in parents)
    labels = ["F30", "Q1", "T27", "其他10个位点"]
    values = [counts["30"], counts["1"], counts["27"], sum(v for k, v in counts.items() if k not in {"30", "1", "27"})]
    fig, axes = plt.subplots(1, 2, figsize=(10.4, 3.8), gridspec_kw={"width_ratios": [1.15, 1]})
    ax = axes[0]
    stages = ["允许单突", "幅度短名单", "风险审查后", "父单突"]
    reviewed_count = single_gate["strict_core_count"] + single_gate["controlled_tradeoff_count"]
    stage_counts = [847, single_gate["magnitude_shortlist_count"], reviewed_count, len(parents)]
    x = range(len(stages))
    ax.plot(x, stage_counts, color=BLUE, linewidth=2.5, marker="o", markersize=8)
    for idx, value in enumerate(stage_counts):
        ax.text(idx, value + 25, str(value), ha="center", fontsize=10, fontweight="bold")
    ax.set_xticks(list(x), stages)
    ax.set_ylabel("候选数")
    ax.set_ylim(0, 930)
    ax.set_title("A  单突连续筛选")
    ax.spines[["top", "right"]].set_visible(False)

    ax = axes[1]
    bars = ax.bar(labels, values, color=[BLUE, BLUE, BLUE, "#9AB6C8"], width=0.62)
    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, value + 0.18, str(value), ha="center", fontsize=11)
    ax.set_ylabel("保留单突数")
    ax.set_ylim(0, 11.5)
    ax.set_title("B  19条父单突的位点配额")
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    parent_plot = fig_dir / "figure_parent19.png"
    fig.savefig(parent_plot, dpi=600, bbox_inches="tight"); plt.close(fig)

    layer_counts = _double_layer_counts(audit)
    fig, axes = plt.subplots(1, 2, figsize=(10.8, 3.9), gridspec_kw={"width_ratios": [1, 1.25]})
    ax = axes[0]
    stages = ["理论组合", "有效双突", "规则合格", "最终选择"]
    numbers = [171, gate["double_candidate_count"], gate["eligible_double_count"], gate["selected_double_count"]]
    x = range(len(stages))
    ax.plot(x, numbers, color=BLUE, linewidth=2.6, marker="o", markersize=9)
    for idx, value in enumerate(numbers):
        ax.text(idx, value + 7, str(value), ha="center", fontsize=11, fontweight="bold")
    ax.set_xticks(list(x), stages)
    ax.set_ylabel("候选数")
    ax.set_ylim(0, 195)
    ax.set_title("A  双突数量漏斗")
    ax.spines[["top", "right"]].set_visible(False)

    ax = axes[1]
    layer_keys = ["A_three_families", "B_two_families", "C_one_family", "ineligible"]
    layer_labels = ["A：3家族", "B：2家族", "C：1家族", "0家族"]
    centers = list(range(len(layer_keys)))
    width = 0.25
    series = [("162条总空间", "total", "#A9C4DD"), ("84条规则合格", "eligible", "#5F96C3"), ("11条终选", "selected", "#174A75")]
    for offset, (name, key, color) in zip((-width, 0, width), series):
        vals = [layer_counts[key][layer] for layer in layer_keys]
        bars = ax.bar([x + offset for x in centers], vals, width=width, label=name, color=color)
        for bar, value in zip(bars, vals):
            if value:
                ax.text(bar.get_x() + bar.get_width() / 2, value + 1.2, str(value), ha="center", fontsize=8)
    ax.set_xticks(centers, layer_labels)
    ax.set_ylabel("双突数")
    ax.set_ylim(0, 90)
    ax.set_title("B  证据层数量")
    ax.legend(frameon=False, fontsize=8)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    double_plot = fig_dir / "figure_double_selection.png"
    fig.savefig(double_plot, dpi=600, bbox_inches="tight"); plt.close(fig)
    sources.update({"parent19": parent_plot, "double": double_plot})
    return sources


def _add_footer(doc: Document) -> None:
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(p.add_run("Nb252 BL21表达量优化项目报告"), size=8, color=GRAY)


def _build_report(
    out: Path,
    parents: list[dict[str, str]],
    doubles: list[dict[str, str]],
    panel: list[dict[str, str]],
    gate: dict,
    figures: dict[str, Path],
    single_gate: dict,
    layer_counts: dict[str, dict[str, int]],
) -> Path:
    doc = Document()
    _style_document(doc)
    p = doc.add_paragraph(style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run("Nb252纳米抗体BL21表达量优化项目报告"), size=21, bold=True)
    p = doc.add_paragraph(style="Subtitle"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run("结构与天然保守性约束下的19条单突和11条双突候选"), size=12)
    doc.add_paragraph()
    _table(doc, ["项目", "当前结果"], [
        ["设计目标", "提高Nb252在BL21体系中的表达相关性质"],
        ["权威亲本", "128 aa；末端SSGS保持不变"],
        ["完整单突空间", "847条允许单突"],
        ["最终候选面板", "19条单突＋11条双突，共30条"],
        ["实验对照", "WT独立设置，不占30条候选名额"],
    ], [1.35, 5.15], 9.4)

    _new_page_heading(doc, "执行摘要")
    _point(doc, "全部候选首先满足实验界面、天然高保守位点、Cys22/Cys95及末端SSGS保护规则。")
    _point(doc, "847条允许单突已完成NetSolP、NanoMelt和AntiFold统一评价，并按预定义幅度档而非同档小数差异筛选。")
    _point(doc, "单突筛选依次经过幅度短名单、序列风险与性质恶化审查以及位点/替换多样性约束，最终保留19条父单突：F30、Q1、T27各3条，其余10个位点各1条。")
    _point(doc, "19条父单突产生171个理论无序对；去除9个同位点互斥组合后，162条有效双突均重新计算完整序列的NetSolP和NanoMelt结果。")
    _point(doc, "84条双突通过无硬风险、无中等或明显性质恶化、且至少一个预测器家族中等或明显有利的资格门。")
    _point(doc, "全局分类优化选择11条双突：A层3条、B层4条、C层4条；最终形成19条单突＋11条双突的30条候选面板。")
    _table(doc, ["阶段", "数量", "含义"], [
        ["允许单突空间", 847, "完整预测空间"],
        ["幅度短名单", single_gate["magnitude_shortlist_count"], "至少一个家族中等/明显有利且无明显不利"],
        ["风险审查后", single_gate["strict_core_count"] + single_gate["controlled_tradeoff_count"], "26条严格核心＋10条受控权衡"],
        ["父单突", 19, "作为独立候选和双突组成项"],
        ["有效双突", 162, "不同位置的全部可实现组合"],
        ["双突资格门通过", 84, "进入全局组合选择"],
        ["最终面板", 30, "19条单突＋11条双突"],
    ], [1.6, 0.8, 4.1], 9.0)

    _new_page_heading(doc, "1. 设计目标与不可变约束")
    doc.add_paragraph("本项目以Nb252在BL21体系中的表达量优化为目标。当前计算不显式优化亲和力，也不使用Rosetta类评分。实验复合物中定义的界面、天然高保守位点、二硫键和末端构建体序列在候选生成前即被冻结。")
    _table(doc, ["约束", "执行规则"], [
        ["实验界面", "24个实验复合物界面位置全部不突变"],
        ["天然保守性", "Nb252自身等于全局/邻域天然优势残基的高置信位置冻结"],
        ["共识回变", "Q5不开放任意扫描，仅允许Q5V天然共识回变"],
        ["构建体", "保持128 aa、Cys22/Cys95及末端SSGS，不引入新Cys"],
        ["组合范围", "双突只能由批准的19条父单突在不同位置组合，不生成三突或更高阶突变"],
    ], [1.45, 5.05], 9.1)
    _figure(doc, figures["constraints"], "图1  Nb252天然保守性与设计约束轨道。", 16.2)

    _new_page_heading(doc, "2. 天然VHH保守性参照")
    doc.add_paragraph("保守性参照来自4,059条公开非冗余天然VHH。经IMGT编号、完整性审核和90% identity去冗余后得到3,784个有效簇；其中Nb252邻域包含1,532个有效簇。邻域用于避免远离Nb252框架背景的VHH主导位置频率。")
    _point(doc, "位置覆盖率不低于0.80，有效簇数不少于50。")
    _point(doc, "邻域优势残基频率不低于0.90，全局优势残基频率不低于0.80，且两者一致。")
    _point(doc, "只有Nb252自身残基等于该优势残基时，位置才因天然保守性完全冻结。")
    _point(doc, "若Nb252偏离高保守共识，仅开放到共识残基的单一回变，并继续服从其他硬约束。")

    doc.add_heading("3. 工具验证与使用边界", level=1)
    doc.add_paragraph("47条BL21产量记录用于判断预测指标的适用边界。连续相关、来源分层分析和序列簇外分类表明，现有工具不能被解释为经过外部验证的mg/L预测器，但可作为相互独立的性质证据与风险约束。")
    _figure(doc, figures["validation"], "图2  多种预测指标与47条BL21产量记录的连续关联和展示性分类结果。", 16.2)
    _table(doc, ["工具", "本项目中的角色", "不作何种解释"], [
        ["NetSolP U/S", "可用性和可溶性软偏好", "不直接换算为mg/L"],
        ["NanoMelt Tm", "避免明显热稳定性退化", "不作为独立产量排序器"],
        ["AntiFold", "结构环境下逐位点氨基酸兼容性", "不等同于表达量或双突上位性"],
    ], [1.2, 2.45, 2.85], 8.8)

    _new_page_heading(doc, "4. 847条允许单突与幅度分档")
    doc.add_paragraph("硬约束合并后，47个常规可扫位置产生846条非Cys替换，加上Q5V共识回变，共847条。NetSolP和NanoMelt覆盖全部候选；AntiFold优先采用实验复合物视图，实验坐标不可评价位置以独立AF3 VHH单体视图补充并保留来源标记。")
    _figure(doc, figures["landscape"], "图3  847条单突的NetSolP ΔU、NetSolP ΔS、NanoMelt ΔTm和AntiFold ΔlogP景观。", 16.3)
    _table(doc, ["指标", "近似中性/轻度/中等/明显分界"], [
        ["NetSolP ΔU", "0.005 / 0.010 / 0.015"],
        ["NetSolP ΔS", "0.010 / 0.020 / 0.030"],
        ["NanoMelt ΔTm", "0.5 / 1.0 / 1.5 °C"],
        ["AntiFold ΔlogP", "0.5 / 1.5 / 3.0"],
    ], [1.7, 4.8], 9.0)
    doc.add_paragraph("原始连续值完整保留用于追溯，但同一幅度档内部的小数差异不参与候选排序，避免微小模型波动主导实验面板。")

    _new_page_heading(doc, "5. 从847条允许单突到19条父单突")
    doc.add_paragraph("19条父单突由完整847条允许空间沿同一套规则连续筛出，不把任一中间集合视为独立实验阶段。首先要求候选无明显不利指标，且NetSolP、NanoMelt或AntiFold至少一个预测器家族达到中等或明显有利，得到40条幅度短名单。随后排除2条新Pro骨架约束风险和2条同时含两个中等不利指标的候选，得到36条可审查候选，其中26条为无中等不利指标的严格核心，10条为仅含一个中等不利指标的受控权衡。")
    doc.add_paragraph("在可审查候选基础上按实验信息量确定19条父单突。F30、Q1和T27各保留3种理化性质互补的替换，用于区分同一位置的替换类型效应；其余10个位点各保留1条，以扩展位置覆盖。具体替换优先依据预测器家族支持数、明显有利家族数、无硬风险/明显恶化、实验复合物AntiFold来源和理化机制互补性；同一幅度档内的小数不参与排序。T99F作为稳定词假设的单独探索项保留，不表述为多工具共同支持。")
    _figure(doc, figures["parent19"], "图4  847条允许单突到19条父单突的连续筛选及位点配额。", 16.2)
    parent_rows = []
    for row in parents:
        mut = _mutation(row)
        parent_rows.append([
            row["parent19_selection_order"], mut, row["region"],
            BAND_LABELS[row["netsolp_s_magnitude_band"]],
            BAND_LABELS[row["nanomelt_tm_magnitude_band"]],
            BAND_LABELS[row["antifold_magnitude_band"]], SINGLE_REASON[mut],
        ])
    _table(doc, ["序", "单突", "区段", "S", "Tm", "AntiFold", "保留依据"], parent_rows,
           [0.35, 0.55, 0.55, 0.7, 0.7, 0.75, 3.25], 7.7)

    _new_page_heading(doc, "6. 162条双突的完整序列评价")
    doc.add_paragraph("19条父单突理论上形成171个无序组合。F30、Q1、T27各自的不同替换不能同时出现在同一序列中，共去除9个同位点互斥组合，得到162条有效双突。")
    _figure(doc, figures["double"], "图5  从理论组合到最终11条双突的数量漏斗。", 15.8)
    _point(doc, "NetSolP U/S和NanoMelt Tm均使用完整双突序列重新计算。")
    _point(doc, "双突U/S与两个单突简单相加只呈中等一致性，因此不能用父单突分数直接拼接。")
    _point(doc, "AntiFold没有进行双突结构重建；仅保留两个组成位置中较弱的单突兼容性档，不能据此声称双突上位性。")
    _point(doc, "每条双突重新检查完整序列风险、稳定词变化和末端/二硫键合同。")
    _table(doc, ["资格门", "要求"], [
        ["序列风险", "硬序列风险数为0"],
        ["不利性质", "U、S、Tm和AntiFold均不得出现中等或明显不利档"],
        ["有利证据", "至少一个预测器家族达到中等或明显有利"],
        ["稳定词", "仅作后置软偏好，不能挽救明显性质恶化"],
    ], [1.45, 5.05], 9.0)

    _new_page_heading(doc, "7. 11条双突的证据分层与全局选择")
    doc.add_paragraph("先按达到中等或明显有利的预测器家族数对全部162条有效双突分层：A层为3个家族支持，B层为2个家族支持，C层为1个家族支持；0家族组没有达到中等幅度的有利证据。分层的原因是三类工具评价对象不同且没有经过验证的统一加权总分：先比较跨家族证据一致性，可避免同类指标重复投票，也避免同档微小数值决定候选。")
    _table(doc, ["证据层", "定义", "162条中总数", "资格门后", "终选数"], [
        ["A层", "NetSolP、NanoMelt、AntiFold三个家族均中等或明显有利", layer_counts["total"]["A_three_families"], layer_counts["eligible"]["A_three_families"], layer_counts["selected"]["A_three_families"]],
        ["B层", "两个家族中等或明显有利", layer_counts["total"]["B_two_families"], layer_counts["eligible"]["B_two_families"], layer_counts["selected"]["B_two_families"]],
        ["C层", "一个家族中等或明显有利", layer_counts["total"]["C_one_family"], layer_counts["eligible"]["C_one_family"], layer_counts["selected"]["C_one_family"]],
        ["0家族", "没有家族达到中等或明显有利", layer_counts["total"]["ineligible"], 0, 0],
    ], [0.65, 3.55, 0.8, 0.8, 0.7], 8.3)
    doc.add_paragraph("资格门进一步排除存在硬序列风险、任一U/S/Tm/AntiFold指标达到中等或明显不利、或没有任何家族达到中等有利的组合。最终84条合格双突由A层5条、B层39条和C层40条组成。A/B/C是证据广度分层，不是要求某一层全部纳入的名单；最终双突名额固定为11条，因此不能把44条合格A/B层组合全部收入。随后采用按优先级逐项固定最优值的分类优化选择整体11条面板，而不是按某个连续总分从高到低截取。")
    _point(doc, "A层5条中，F30S+Q5V、F30A+Q5V和F30R+Q5V共享同一位置对；为避免只检验同一组合假设，每个位置对最多保留1条。因此在其余约束同时成立时，A层可实现的最大数量是3条，而不是人为只取3条。")
    _point(doc, "固定A层最优数量后，再最大化B层数量。B层候选集中在已经高频出现的组成突变和位置；在每个精确组成突变最多使用2次、每个原序列位置最多使用3次的限制下，B层可实现的最大数量是4条。")
    _point(doc, "剩余4个名额由同样通过资格门的C层候选补足。它们不是用来替代更高层的可行候选，而是在A=3、B=4的约束最优值已经固定后，扩展到13种组成单突和10个原序列位置，提高实验信息量并降低少数位点或单突反复出现的风险。")
    _point(doc, "随后减少软序列风险和中等/明显不利的非加和残差。")
    _point(doc, "最后增加不同组成单突和不同位置覆盖；每个位置对最多1条、每个确切单突最多使用2次、每个位置最多出现3次。")
    _point(doc, "11条双突覆盖13种组成单突和10个原序列位置；没有双突稳定词净新增，稳定词假设由单突T99F保留。")
    double_rows = []
    for row in doubles:
        residual = int(row["moderate_or_strong_adverse_interaction_residual_count"])
        note = f"{row['favorable_family_count']}个家族支持"
        if residual:
            note += f"；{residual}项不利非加和残差，但双突实际值仍通过硬门"
        else:
            note += "；无中等/明显不利残差"
        double_rows.append([
            row["double_selection_order"], row["mutation_set"], row["evidence_layer"][0],
            _fmt(row["netsolp_u_delta_vs_wt"]), _fmt(row["netsolp_s_delta_vs_wt"]),
            _fmt(row["nanomelt_tm_c_delta_vs_wt"], 2), BAND_LABELS[row["antifold_worst_component_band"]], note,
        ])
    _table(doc, ["序", "双突", "层", "ΔU", "ΔS", "ΔTm °C", "AntiFold", "入选说明"], double_rows,
           [0.3, 0.9, 0.3, 0.55, 0.55, 0.65, 0.8, 2.75], 7.2)

    _new_page_heading(doc, "8. 最终30条候选面板")
    doc.add_paragraph("最终面板包含19条单突和11条双突。单突用于直接估计组成替换的实验效应；双突用于检验两个受支持替换在同一完整序列中的累积、拮抗或非加和表现。WT作为独立实验对照。")
    _figure(doc, figures["final"], "图6  最终19条单突和11条双突的分类证据、选择漏斗与位置覆盖。", 16.4)
    _table(doc, ["面板组成", "数量", "主要用途"], [
        ["父单突", 19, "测量每个组成替换的独立表达效应"],
        ["A层双突", 3, "检验三家族一致支持的组合"],
        ["B层双突", 4, "检验双家族支持的组合"],
        ["C层双突", 4, "补充位置与机制多样性"],
        ["WT", 1, "独立对照，不计入30条候选"],
    ], [1.35, 0.7, 4.45], 9.0)

    _new_page_heading(doc, "9. 解释边界与实验建议")
    _point(doc, "所有30条均为计算优先候选，不代表表达量已经实验提高。")
    _point(doc, "双突AntiFold证据来自两个组成位点的保守汇总，没有直接评价双突结构上位性。")
    _point(doc, "F30S+S55G和Q1H+F29T各存在1项中等或明显不利的非加和残差；其双突实际性质仍未达到排除档，应作为实验判读重点。")
    _point(doc, "建议WT与30条候选采用相同载体、信号肽、标签、培养、诱导、裂解、纯化和定量流程。")
    _point(doc, "建议至少设置3个独立生物重复，以纯化后mg/L为主要终点，并同步记录总表达、可溶比例和纯化回收。")
    _point(doc, "双突结果应与两个对应单突同时比较，以区分简单累积、非加和改善和组合拮抗。")

    _new_page_heading(doc, "附录A  19条父单突索引")
    single_panel_rows = []
    double_panel_rows = []
    for row in panel:
        layer, reason = _panel_reason(row)
        values = [row["final_panel_order"], row["mutation_set"], row["regions"], layer, reason]
        if row["candidate_kind"] == "single_mutant":
            single_panel_rows.append(values)
        else:
            double_panel_rows.append(values)
    _table(doc, ["序", "单突", "区段", "证据层", "主要依据"], single_panel_rows,
           [0.35, 0.8, 0.8, 1.0, 3.55], 8.0)
    _new_page_heading(doc, "附录B  11条双突索引")
    _table(doc, ["序", "双突", "区段", "证据层", "主要依据"], double_panel_rows,
           [0.35, 0.9, 0.9, 0.8, 3.55], 8.0)
    doc.add_paragraph("完整128-aa序列及原始机器可读字段随交付CSV和FASTA提供。")

    _add_footer(doc)
    doc.core_properties.title = "Nb252纳米抗体BL21表达量优化项目报告"
    doc.core_properties.subject = "19条单突和11条双突最终候选面板"
    doc.core_properties.author = "Antibody_optimization project"
    report = out / "Nb252_BL21_expression_optimization_project_report.docx"
    doc.save(report)
    return report


def _write_delivery(out: Path, report: Path, panel: list[dict[str, str]], parents: list[dict[str, str]], doubles: list[dict[str, str]]) -> None:
    delivery = out / "delivery"
    delivery.mkdir(parents=True)
    shutil.copy2(report, delivery / report.name)
    shutil.copy2(FINAL_DIR / "nb252_final_30_candidate_panel.csv", delivery / "Nb252_final_30_candidates.csv")
    shutil.copy2(FINAL_DIR / "nb252_final_30_candidate_panel.fasta", delivery / "Nb252_final_30_candidates.fasta")
    shutil.copy2(PARENT_DIR / "expression_single_mutant_parent19.csv", delivery / "Nb252_parent19_single_mutants.csv")
    shutil.copy2(FINAL_DIR / "expression_double_mutant_selected11.csv", delivery / "Nb252_selected11_double_mutants.csv")
    _write_utf8_lf(
        delivery / "README.md",
        "# Nb252 BL21表达量优化交付包\n\n"
        "本目录包含项目报告、最终30条候选CSV/FASTA、19条父单突证据表和11条双突证据表。\n\n"
        "最终候选由19条单突和11条双突组成；WT应作为独立实验对照。所有序列均为128 aa并保留末端SSGS。"
        "预测值用于实验优先级与风险控制，不是实测表达量、溶解度或Tm。\n",
    )


def main() -> int:
    parser = argparse.ArgumentParser(description="Build the current Nb252 19-single plus 11-double report package")
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    out = args.output_dir.resolve()
    if out.exists():
        raise FileExistsError(f"Output directory already exists: {out}")
    out.mkdir(parents=True)
    parents = _csv(PARENT_DIR / "expression_single_mutant_parent19.csv")
    doubles = _csv(FINAL_DIR / "expression_double_mutant_selected11.csv")
    panel = _csv(FINAL_DIR / "nb252_final_30_candidate_panel.csv")
    audit = _csv(FINAL_DIR / "expression_double_mutant_final_selection_audit.csv")
    gate = _json(FINAL_DIR / "expression_final_panel_gate.json")
    single_gate = _json(TRIAL_DIR / "expression_single_mutant_selection_gate.json")
    layer_counts = _double_layer_counts(audit)
    expected_layers = {
        "total": {"A_three_families": 5, "B_two_families": 48, "C_one_family": 80, "ineligible": 29},
        "eligible": {"A_three_families": 5, "B_two_families": 39, "C_one_family": 40, "ineligible": 0},
        "selected": {"A_three_families": 3, "B_two_families": 4, "C_one_family": 4, "ineligible": 0},
    }
    if (
        (len(parents), len(doubles), len(panel), len(audit)) != (19, 11, 30, 162)
        or gate["status"] != "pass"
        or layer_counts != expected_layers
    ):
        raise ValueError("Final 19+11 panel inputs are incomplete or unreleased")
    figures = _make_figures(out, parents, doubles, gate, audit, single_gate)
    report = _build_report(out, parents, doubles, panel, gate, figures, single_gate, layer_counts)
    _write_delivery(out, report, panel, parents, doubles)
    manifest = {
        "schema_version": 2,
        "active_route": "BL21_expression_19_single_plus_11_double",
        "final_candidate_count": 30,
        "single_mutant_count": 19,
        "double_mutant_count": 11,
        "single_mutant_screen_counts": {
            "allowed": 847,
            "magnitude_shortlist": single_gate["magnitude_shortlist_count"],
            "risk_reviewed": single_gate["strict_core_count"] + single_gate["controlled_tradeoff_count"],
            "selected_parent": 19,
        },
        "eligible_double_count": 84,
        "source_double_candidate_count": 162,
        "double_evidence_layer_counts": layer_counts,
        "report_docx": report.name,
        "presentation_pptx": "Nb252_BL21_expression_optimization_presentation.pptx",
        "delivery_directory": "delivery",
    }
    _write_utf8_lf(out / "report_manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2) + "\n")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
