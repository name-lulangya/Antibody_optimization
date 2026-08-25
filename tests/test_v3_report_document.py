from __future__ import annotations

import hashlib
import json
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts/reporting/build_nb252_expression_v3_report.py"
TEMPLATE = (
    ROOT
    / "docs/result_artifacts/weekly_report_result/"
    "report_2026_W34_nb252_expression_route/"
    "Nb252_BL21_expression_optimization_project_report.docx"
)

EXPECTED_SINGLES = [
    "L11Y",
    "F30S",
    "K86S",
    "A23R",
    "Q5V",
    "S55G",
    "K75A",
    "F29Q",
    "K43A",
    "N76G",
    "F30N",
    "K75E",
    "L11M",
    "Q1D",
    "T99F",
]

EXPECTED_DOUBLES = [
    "F30S;Q5V",
    "S55G;K43A",
    "K86S;Q5V",
    "L11Y;K86S",
    "S55G;F30N",
    "N76G;L11M",
    "F30S;K75E",
    "K86S;K43A",
    "A23R;S55G",
    "K43A;N76G",
    "K75E;Q1D",
    "L11Y;K75A",
    "Q5V;N76G",
    "L11Y;Q1D",
    "F30S;Q1D",
]


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _all_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def _table_with_header(
    document: Document,
    first: str,
    second: str,
    third: str | None = None,
):
    matches = []
    for table in document.tables:
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if (
            len(header) >= 2
            and header[0] == first
            and header[1] == second
            and (third is None or (len(header) >= 3 and header[2] == third))
        ):
            matches.append(table)
    assert len(matches) == 1
    return matches[0]


@pytest.fixture(scope="module")
def built_v3_report(tmp_path_factory: pytest.TempPathFactory):
    work = tmp_path_factory.mktemp("v3_report_document")
    output = work / "Nb252_V3_test.docx"
    figures = work / "figures"
    manifest = work / "manifest.json"
    template_bytes = TEMPLATE.read_bytes()

    completed = subprocess.run(
        [
            sys.executable,
            str(SCRIPT),
            "--template",
            str(TEMPLATE),
            "--output",
            str(output),
            "--figures",
            str(figures),
            "--manifest",
            str(manifest),
        ],
        cwd=ROOT,
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    assert completed.returncode == 0, completed.stdout + completed.stderr
    assert output.is_file() and output.stat().st_size > 100_000
    assert manifest.is_file()
    assert TEMPLATE.read_bytes() == template_bytes
    return {
        "document": Document(str(output)),
        "manifest": json.loads(manifest.read_text(encoding="utf-8")),
        "output": output,
        "template_hash": hashlib.sha256(template_bytes).hexdigest(),
    }


def test_real_v3_build_preserves_template_and_manifest_scope(built_v3_report):
    manifest = built_v3_report["manifest"]
    document = built_v3_report["document"]
    assert manifest["historical_template_sha256"] == built_v3_report["template_hash"]
    assert _sha256(TEMPLATE) == built_v3_report["template_hash"]
    assert manifest["document"]["candidate_count"] == 30
    assert manifest["document"]["single_count"] == 15
    assert manifest["document"]["double_count"] == 15
    assert document.core_properties.title == "Nb252纳米抗体BL21表达量优化项目报告（V3）"
    assert document.core_properties.subject == "15条单突和15条双突最终候选面板"
    assert not Path(manifest["document"]["output_docx"]).is_absolute()
    assert all(
        not Path(path).is_absolute() for path in manifest["source_artifacts"].values()
    )
    assert all(
        not Path(value).is_absolute()
        for metadata in manifest["figures"].values()
        for key, value in metadata.items()
        if key in {"path", "png", "svg"}
    )
    assert manifest["scope"] == {
        "ppt_created": False,
        "delivery_archive_created": False,
        "historical_v2_assets_modified": False,
        "antifold_role": "negative_risk_exclusion_only",
    }


def test_report_contains_exact_released_15_plus_15_identities(built_v3_report):
    document = built_v3_report["document"]
    single_table = _table_with_header(document, "展示序", "突变", "完整序列（128 aa）")
    double_table = _table_with_header(document, "展示序", "组合", "完整序列（128 aa）")

    assert [row.cells[1].text.strip() for row in single_table.rows[1:]] == EXPECTED_SINGLES
    assert [row.cells[1].text.strip() for row in double_table.rows[1:]] == EXPECTED_DOUBLES
    assert len(single_table.rows) == 16
    assert len(double_table.rows) == 16


def test_report_excludes_historical_v2_panel_semantics(built_v3_report):
    text = _all_text(built_v3_report["document"])
    forbidden = (
        "19条父单突",
        "19条单突",
        "11条双突",
        "19单突＋11双突",
        "19+11",
        "162条双突",
        "162个双突",
    )
    assert not [phrase for phrase in forbidden if phrase in text]


def test_antifold_language_is_exclusion_only_and_never_positive_credit(built_v3_report):
    text = _all_text(built_v3_report["document"])
    assert "AntiFold仅用于风险排除" in text
    assert "不提议" in text
    assert "不奖励" in text
    assert "不排序" in text
    assert (
        "双突不计算AntiFold分数" in text
        or "双突不运行AntiFold联合评分" in text
    )
    assert "AntiFold改善从不作为双突入选理由" in text
    assert "不表示正向支持" in text


def test_collaborator_facing_terms_and_table_structure_are_explicit(built_v3_report):
    document = built_v3_report["document"]
    text = _all_text(document)

    assert "“硬约束”指禁止改变24个实验直接界面位点" in text
    assert "“序列硬风险”指新增Pro造成主链构象约束" in text
    assert "7-aa窗口出现至少6个疏水残基" in text
    assert "净侧链电荷绝对值达到至少5" in text
    assert "“父单突”是指获准作为双突组成部分的单突" in text
    assert "稳定词是用简并氨基酸符号表示" in text

    excluded = _table_with_header(document, "排除突变", "主要物理风险", "专家排除依据")
    assert [row.cells[0].text.strip() for row in excluded.rows[1:]] == [
        "A49F",
        "A49M",
        "S50F",
        "R71G",
        "A96R",
    ]

    parents = _table_with_header(document, "序", "单突", "区域")
    doubles = _table_with_header(document, "序", "双突", "区域")
    assert len(parents.rows) == 16
    assert len(doubles.rows) == 16
    assert parents.rows[0].cells[6].text.strip() == "结构证据/判断"
    assert doubles.rows[0].cells[6].text.strip() == "结构证据/风险"

    forbidden_terms = (
        "实验验证建议",
        "结构/责任",
        "化学责任",
        "序列责任",
        "新增脱酰胺责任",
        "实验QC",
        "功能QC",
        "7.1 展示顺序",
        "7.2 展示顺序",
        "9.1 展示顺序",
        "9.2 展示顺序",
    )
    assert not [term for term in forbidden_terms if term in text]


def test_every_report_table_is_strict_unshaded_three_line_table(built_v3_report):
    document = built_v3_report["document"]
    assert len(document.tables) >= 10
    for table in document.tables:
        tbl_pr = table._tbl.tblPr
        borders = tbl_pr.find(qn("w:tblBorders"))
        assert borders is not None
        for edge in ("top", "bottom"):
            node = borders.find(qn(f"w:{edge}"))
            assert node is not None and node.get(qn("w:val")) == "single"
        for edge in ("left", "right", "insideH", "insideV"):
            node = borders.find(qn(f"w:{edge}"))
            assert node is not None and node.get(qn("w:val")) == "nil"

        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            assert tc_borders is not None
            bottom = tc_borders.find(qn("w:bottom"))
            assert bottom is not None and bottom.get(qn("w:val")) == "single"

        for row in table.rows:
            for cell in row.cells:
                assert cell._tc.get_or_add_tcPr().find(qn("w:shd")) is None
