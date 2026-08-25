"""Build the audience-facing Nb252 V3 project report from validated V3 data.

The builder deliberately reuses only the visual template of the historical report.
Every scientific fact comes from ``v3_report_data``.  AntiFold is rendered only as a
negative risk-exclusion rule: it never proposes, rewards, or ranks a candidate, and
there is no double-mutant AntiFold score.
"""

from __future__ import annotations

import hashlib
import json
import os
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips


FONT = "Microsoft YaHei"
MONO_FONT = "Consolas"
BLACK = "000000"
GRAY = "666666"
BLUE = "2F6F9F"

BAND_LABELS = {
    "strong_favorable": "显著有利",
    "moderate_favorable": "中等有利",
    "weak_favorable": "轻度有利",
    "negligible": "近似中性",
    "weak_adverse": "轻度不利",
    "moderate_adverse": "中等不利",
    "strong_adverse": "显著不利",
}

EXPERT_LABELS = {
    "reasonable": "结构上合理",
    "reasonable_with_caution": "谨慎合理",
    "structurally_concerning": "存在结构担忧",
    "indeterminate": "暂无法确定",
    "favorable": "预期有利",
    "neutral_or_uncertain": "中性或不确定",
    "unfavorable": "预期不利",
    "high": "高",
    "medium": "中",
    "low": "低",
}


class V3ReportDocumentError(ValueError):
    """Raised when report data cannot be rendered without semantic ambiguity."""


def _font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: str = BLACK,
    name: str = FONT,
) -> None:
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)
    run.font.color.rgb = RGBColor.from_string(color)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _clear_template_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(0.9)
    specs = {
        "Normal": (10.5, False, 0, 5, 1.25),
        "Title": (21, True, 0, 6, 1.0),
        "Subtitle": (12, False, 0, 5, 1.0),
        "Heading 1": (15, True, 8, 7, 1.05),
        "Heading 2": (12.5, True, 6, 5, 1.05),
        "Caption": (9, False, 3, 6, 1.0),
        "List Bullet": (10.5, False, 0, 4, 1.2),
    }
    for name, (size, bold, before, after, spacing) in specs.items():
        style = doc.styles[name]
        style.font.name = FONT
        fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{key}"), FONT)
        style.font.color.rgb = RGBColor.from_string(BLACK)
        style.font.size = Pt(size)
        style.font.bold = bold
        fmt = style.paragraph_format
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.line_spacing = spacing
    bullet = doc.styles["List Bullet"].paragraph_format
    bullet.left_indent = Cm(0.65)
    bullet.first_line_indent = Cm(-0.32)


def _page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


def _set_footer(doc: Document) -> None:
    for section in doc.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(paragraph.add_run("Nb252 BL21表达优化项目报告（V3）  ·  "), size=8, color=GRAY)
        _page_field(paragraph)


def _paragraph(doc: Document, text: str, *, bold_lead: str | None = None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        _font(paragraph.add_run(bold_lead), bold=True)
        _font(paragraph.add_run(text[len(bold_lead) :]))
    else:
        _font(paragraph.add_run(text))
    return paragraph


def _bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    _font(paragraph.add_run(text))


def _new_section(doc: Document, title: str, *, page_break_before: bool = True) -> None:
    heading = doc.add_heading(title, level=1)
    heading.paragraph_format.page_break_before = page_break_before


def _heading(doc: Document, title: str, *, page_break_before: bool = False) -> None:
    heading = doc.add_heading(title, level=2)
    heading.paragraph_format.page_break_before = page_break_before


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    tr_pr.append(marker)


def _border(parent, edge: str, value: str, size: int = 0) -> None:
    node = parent.find(qn(f"w:{edge}"))
    if node is None:
        node = OxmlElement(f"w:{edge}")
        parent.append(node)
    node.set(qn("w:val"), value)
    node.set(qn("w:sz"), str(size))
    node.set(qn("w:color"), BLACK)
    node.set(qn("w:space"), "0")


def _cell(cell, value: object, *, bold: bool, size: float, mono: bool = False) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.03
    _font(
        paragraph.add_run(str(value)),
        size=size,
        bold=bold,
        name=MONO_FONT if mono else FONT,
    )
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, amount in (("top", 70), ("bottom", 70), ("start", 95), ("end", 95)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(amount))
        node.set(qn("w:type"), "dxa")


def _table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
    weights: Sequence[float],
    *,
    size: float = 8.1,
    mono_columns: Iterable[int] = (),
):
    if not (len(headers) == len(weights) and all(len(row) == len(headers) for row in rows)):
        raise V3ReportDocumentError("Table shape does not match its declared columns")
    mono_columns = set(mono_columns)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _repeat_header(table.rows[0])
    _cant_split(table.rows[0])
    for idx, value in enumerate(headers):
        _cell(table.rows[0].cells[idx], value, bold=True, size=size)
    for values in rows:
        row = table.add_row()
        _cant_split(row)
        for idx, value in enumerate(values):
            _cell(row.cells[idx], value, bold=False, size=size, mono=idx in mono_columns)
    total = 9500
    widths = [round(total * value / sum(weights)) for value in weights]
    widths[-1] += total - sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Twips(width)
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    _border(borders, "top", "single", 12)
    _border(borders, "bottom", "single", 12)
    for edge in ("left", "right", "insideH", "insideV"):
        _border(borders, edge, "nil")
    for cell in table.rows[0].cells:
        tc_borders = cell._tc.get_or_add_tcPr().find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            cell._tc.get_or_add_tcPr().append(tc_borders)
        _border(tc_borders, "bottom", "single", 8)
    return table


def _figure(doc: Document, path: Path, caption: str, *, width_cm: float = 16.4) -> None:
    if not path.is_file():
        raise V3ReportDocumentError(f"Report figure is missing: {path}")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(cap.add_run(caption), size=9)


def _mutation(row: Mapping[str, Any]) -> str:
    label = str(row.get("mutation_reported_label", ""))
    return label.rsplit(" ", 1)[-1] if label else str(row.get("mutation_set", ""))


def _band(value: str) -> str:
    return BAND_LABELS.get(value, value)


def _metric(value: str, band: str, *, tm: bool = False) -> str:
    suffix = " °C" if tm else ""
    return f"{float(value):+.3f}{suffix}\n{_band(band)}"


def _structure_source(value: str) -> str:
    if "af3" in value.lower():
        return "AF3补充上下文"
    return "实验复合物"


def _soft_risk(value: str) -> str:
    if value == "new_deamidation_motif":
        return "新增NG脱酰胺基序"
    if value == "more_M_or_W":
        return "增加Met氧化敏感性"
    return "未记录额外化学风险"


def _parent_note(row: Mapping[str, Any]) -> str:
    mutation = _mutation(row)
    mapping = {
        "L11Y": "S显著、U中等改善；仅有AF3局部坐标，结构判断置信度较低。",
        "F30S": "S显著、U中等改善；表面去疏水，但需关注CDR1预组织。",
        "K86S": "Tm显著、S中等改善；外露Lys→Ser并保留极性。",
        "A23R": "U和Tm均中等改善；靠近Cys22，大体积带电侧链需实验关注。",
        "Q5V": "Tm显著改善且为唯一允许的天然共识回变；表面疏水性可能上升。",
        "S55G": "Tm显著改善；Gly可能增加CDR2柔性，作为机制权衡候选。",
        "K75A": "Tm显著改善；结构易容纳，但U轻度下降。",
        "F29Q": "S中等改善；实验坐标缺失，保留为CDR1极性化假设。",
        "K43A": "S中等、Tm轻度改善；表面电荷调整假设。",
        "N76G": "Tm中等、U轻度改善；正φ构象支持Gly转角。",
        "F30N": "U显著、S中等改善；新增NG脱酰胺基序，明确标注化学风险。",
        "K75E": "Tm显著改善；电荷反转效应依赖局部pH与盐环境。",
        "L11M": "U显著改善；仅AF3坐标且增加Met氧化敏感性。",
        "Q1D": "S中等、U轻度改善；外露N端负电荷水化假设。",
        "T99F": "稳定词新增的探索对照；U无改善、S/Tm轻度不利并有中等结构担忧。",
    }
    return mapping[mutation]


def _double_note(row: Mapping[str, Any]) -> str:
    mutation = str(row["mutation_set"])
    mapping = {
        "F30S;Q5V": "CDR1去疏水与FR1共识回变互补，三项性质均达到中等或显著有利。",
        "S55G;K43A": "CDR2构象假设与FR2表面电荷调整互补，三项性质均有支持。",
        "K86S;Q5V": "两处远距替换兼顾表面极性与共识回变，Tm改善最明显。",
        "L11Y;K86S": "三项性质均有支持；L11仅有AF3坐标，结构证据置信度较低。",
        "S55G;F30N": "三项均显著有利；保留为高收益—新增NG脱酰胺风险的明确权衡。",
        "N76G;L11M": "转角Gly与FR1替换互补；L11仅AF3坐标并增加Met氧化敏感性。",
        "F30S;K75E": "S与Tm均显著有利，覆盖CDR1去疏水与FR3电荷反转两类机制。",
        "K86S;K43A": "两个远距外露框架位点的电荷/极性调整，S与Tm均显著有利。",
        "A23R;S55G": "U中等、Tm显著有利；保留reported 23机制并标注Cys22邻近担忧。",
        "K43A;N76G": "表面电荷调整与转角构象假设互补，S中等、Tm显著有利。",
        "K75E;Q1D": "两个远距电荷假设，S中等、Tm显著有利；电荷环境仍需实验验证。",
        "L11Y;K75A": "U中等、Tm显著有利；L11仅AF3坐标，作为低置信度结构假设。",
        "Q5V;N76G": "共识回变与Gly转角互补，S中等、Tm显著有利。",
        "L11Y;Q1D": "两处水化假设使U/S有支持；L11仅AF3坐标，Tm近似中性。",
        "F30S;Q1D": "CDR1去疏水与N端负电荷互补，U/S均中等有利，Tm近似中性。",
    }
    return mapping[mutation]


def _add_cover(doc: Document, data: Mapping[str, Any]) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(title.add_run("Nb252 BL21表达优化项目报告（V3）"), size=21, bold=True)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(subtitle.add_run("结构与天然保守性约束下的15条单突与15条双突候选"), size=12)
    doc.add_paragraph()
    _table(
        doc,
        ["项目要素", "当前定义"],
        [
            ["优化目标", "Nb252在BL21体系中的表达产量"],
            ["权威母本", f"128 aa完整构建；末端SSGS保持不变"],
            ["最终计算面板", "15条单突＋15条双突；WT另设对照"],
            ["突变编号", "均为128-aa reported-sequence位置"],
            ["结果性质", "计算优先实验假设，尚未由Nb252突变体表达实验验证"],
        ],
        [1, 3],
        size=9.3,
    )


def _add_executive_summary(doc: Document, data: Mapping[str, Any]) -> None:
    _new_section(doc, "执行摘要")
    _bullet(doc, "权威母本为完整128-aa Nb252构建，末端SSGS属于构建设计并保持不变。所有候选均使用reported-sequence编号。")
    _bullet(doc, "候选生成首先冻结实验复合物定义的24个直接界面位点、天然保守位点、Cys22/Cys95及末端125–128位；Q5只允许Q5V天然共识回变。")
    _bullet(doc, "在847条允许单突中，NetSolP U、NetSolP S与NanoMelt预测Tm按预设幅度档分别评价；档内微小小数差异不用于排序。")
    _bullet(doc, "AntiFold仅用于风险排除：它不提议、不奖励、也不排序候选。单突仅在ΔlogP≤−3且处于该位点20种状态最差4名时排除；双突不计算AntiFold分数。")
    _bullet(doc, "最终选定15条用于构建双突的组成单突（下文简称“父单突”）；完整枚举105个理论配对并去除3个同位点无效组合后，对102条双突全部重算U/S/Tm，再选定15条双突。")
    _bullet(doc, "最终30条均为计算优先实验假设。直接界面未突变并不等同于结合保持已被证明，仍需通过表达、可溶性、单体状态与结合/功能实验验证。")
    _heading(doc, "内容概览")
    _table(
        doc,
        ["部分", "核心问题"],
        [
            ["1–3", "母本、硬约束与天然保守性如何定义"],
            ["4–5", "预测工具为何保留、幅度档如何使用"],
            ["6–7", "847条单突如何收缩为15条父单突"],
            ["8–9", "102条双突如何完整评价并选出15条"],
            ["10–11", "最终面板与风险边界"],
            ["附录", "指标定义、完整序列与证据符号"],
        ],
        [1, 3],
        size=9.2,
    )


def _add_parent_and_constraints(doc: Document, data: Mapping[str, Any]) -> None:
    _new_section(doc, "1. 项目目标与母本身份")
    _paragraph(doc, "本项目以Nb252在BL21体系中的表达产量为优化目标。当前路线不显式优化亲和力；实验复合物中的直接界面残基被完整冻结，以降低直接破坏已观察结合界面的风险。")
    _paragraph(doc, "权威母本是完整128-aa reported sequence。C端SSGS属于构建本身，不是可删除的连接肽；reported 125–128位在全部候选中保持不变。报告中的L11Y、F30S等标签均指reported-sequence位置，不能直接替换为IMGT编号。")
    _table(
        doc,
        ["身份项目", "冻结定义"],
        [
            ["母本长度", "128 aa"],
            ["二硫键半胱氨酸", "reported C22与C95"],
            ["不可变末端", "reported 125–128：SSGS"],
            ["最终候选", "15条单突＋15条双突；WT不计入30条"],
        ],
        [1, 2.8],
        size=9.2,
    )

    _new_section(doc, "2. 硬约束与结构证据边界")
    constraints = data["constraints"]
    _paragraph(doc, f"候选生成前共冻结{constraints['hard_frozen_position_count']}个reported位置，其中包括实验复合物按重原子中心距离严格<4.0 Å定义的{constraints['experimental_interface_position_count']}个直接界面位点、C22/C95、天然保守位点和末端SSGS。约束在性质评分前执行，因此被保护位置不会进入候选空间。")
    _bullet(doc, "实验NK2R–Nb252复合物是结合姿态与直接界面的主要结构证据。")
    _bullet(doc, "实验结构缺失的局部坐标仍视为“不可评价”，AF3仅提供独立的预测上下文，不能改写为实验观察。")
    _bullet(doc, "冻结直接界面只降低直接接触残基被改变的风险；非界面CDR或界面邻近位点仍可能通过构象产生间接影响。")
    _bullet(doc, "双突阶段没有进行双侧链重建、能量最小化或双突AntiFold评分；结构部分属于WT几何分层和组成单突专家证据复核。")
    _table(
        doc,
        ["约束类型", "执行方式", "可解释边界"],
        [
            ["实验直接界面", "24个位置全部禁止突变", "不是能量热点定义，也不证明非界面突变必然保留结合"],
            ["天然保守性", "母本等于天然优势残基时冻结", "非共识高保守位置只开放共识回变"],
            ["二硫键", "C22/C95禁止突变且禁止新增Cys", "不替代折叠与二硫键的实验检查"],
            ["末端构建", "125–128位SSGS保持不变", "NanoMelt评分域为前126 aa，见方法说明"],
        ],
        [1.1, 1.5, 2.3],
        size=8.7,
    )


def _add_conservation(doc: Document, data: Mapping[str, Any], figures: Mapping[str, Path]) -> None:
    _new_section(doc, "3. 天然VHH保守性约束")
    conservation = data["conservation"]
    _paragraph(doc, f"天然保守性使用4,059条公开非冗余天然VHH记录，其中4,057条通过编号与序列质量门。按90%序列一致性去冗余后形成3,784个有效簇；“有效簇”表示高度相似的序列只贡献一次权重，避免同一谱系的重复序列放大频率。")
    _paragraph(doc, f"为更贴近Nb252的序列背景，另定义包含1,564条邻近VHH、1,532个有效簇的局部邻域。全局VHH提供广泛天然约束，邻域VHH提供与Nb252框架和CDR背景更接近的局部共识；两者需方向一致才用于强冻结。")
    _bullet(doc, "要求邻域覆盖率≥0.80、邻域优势残基频率≥0.90、全局优势残基频率≥0.80，且全局与邻域优势残基一致。")
    _bullet(doc, "只有当Nb252自身残基等于共同优势残基时，才能因天然保守性完全冻结。")
    _bullet(doc, "若Nb252位于高保守但非共识状态，则只开放母本→共同优势残基的回变；本项目仅Q5V符合该情况。")
    _figure(
        doc,
        Path(data["figures"]["conservation"]),
        "图1  Nb252天然保守性与设计约束轨迹。保守性依据全局VHH与Nb252邻域VHH的去冗余簇频率共同定义；阴影/标记对应reported-sequence位置。",
        width_cm=16.4,
    )


def _add_tools(doc: Document, data: Mapping[str, Any], figures: Mapping[str, Path]) -> None:
    _new_section(doc, "4. 预测工具验证与在V3中的角色")
    tools = data["tool_validation"]
    _paragraph(doc, "工具与47条已有产量序列的关系同时从连续关联和固定5 mg/L展示阈值下的分类表现观察。LTT/WCC的近似个体产量保留为数值；LLJ的组别或下界不被伪造为个体点值。来源分层Spearman用于减弱来源间基线差异，逐样本留一与序列簇留一用于评估对未见样本和未见近邻序列的外推。")
    _figure(
        doc,
        Path(figures["tool_validation_summary"]),
        "图2  表达相关工具的连续关联与分类表现。5 mg/L仅用于展示；性能不足以将任一工具解释为已验证的Nb252产量预测器。AntiFold不参与正向推荐。",
        width_cm=14.7,
    )
    _table(
        doc,
        ["指标/工具", "含义", "V3中的角色"],
        [
            ["NetSolP U", "0–1综合可用性预测分", "与S分开评价的性质证据；不是mg/L预测值"],
            ["NetSolP S", "0–1预测溶解性分", "与U分开评价；二者来自同一模型，不视为独立模型投票"],
            ["NanoMelt Tm", "预测表观熔解温度（°C）", "热稳定性方向约束；不作为产量排序器"],
            ["AntiFold", "给定结构背景下的残基条件相容性", "仅作风险排除；不提议、不奖励、不排序候选"],
            ["RP3Net / PLM_Sol", "额外表达/溶解度预测器", "验证后不进入V3候选排名，避免同类工具重复投票"],
        ],
        [1.05, 1.55, 2.25],
        size=8.6,
    )
    _heading(doc, "AntiFold双条件排除规则")
    _paragraph(doc, "AntiFold仅在单突ΔlogP≤−3且该突变处于同一位置20种氨基酸状态的最差4名时排除；两条件必须同时满足。未被排除只表示未触发预设风险规则，不表示AntiFold支持或推荐。双突不运行AntiFold联合评分，也不相加组成单突数值；两个组成单突均须通过该风险门，AntiFold改善从不作为双突入选理由。")


def _add_magnitude(doc: Document, data: Mapping[str, Any]) -> None:
    _new_section(doc, "5. 847条允许单突与幅度分档")
    landscape = data["single_landscape"]
    _paragraph(doc, f"硬约束后形成{landscape['allowed_candidate_count']}条允许单突，覆盖{landscape['reported_position_count']}个reported位置。该空间由47个常规开放位点的非WT、非Cys替换，加上Q5V这一条受限共识回变构成。所有847条均完成NetSolP U/S、NanoMelt Tm、AntiFold风险门与序列风险检查。")
    _paragraph(doc, f"AntiFold双条件风险门排除{landscape['antifold_veto_count']}条，{landscape['antifold_veto_pass_count']}条未触发；其中{landscape['qualified_count']}条符合预设的U/S/Tm性质条件。AntiFold的151条排除只代表结构相容性风险门，不改变U/S/Tm的性质幅度定义。")
    thresholds = data["magnitude_thresholds"]
    _table(
        doc,
        ["性质", "近似中性/轻度边界", "轻度/中等边界", "中等/显著边界", "方向"],
        [
            ["NetSolP U", f"|Δ|={thresholds['netsolp_u']['negligible_to_weak_absolute_boundary']:.3f}", f"|Δ|={thresholds['netsolp_u']['weak_to_moderate_absolute_boundary']:.3f}", f"|Δ|={thresholds['netsolp_u']['moderate_to_strong_absolute_boundary']:.3f}", "Δ>0有利"],
            ["NetSolP S", f"|Δ|={thresholds['netsolp_s']['negligible_to_weak_absolute_boundary']:.3f}", f"|Δ|={thresholds['netsolp_s']['weak_to_moderate_absolute_boundary']:.3f}", f"|Δ|={thresholds['netsolp_s']['moderate_to_strong_absolute_boundary']:.3f}", "Δ>0有利"],
            ["NanoMelt Tm", f"|Δ|={thresholds['nanomelt_tm_c']['negligible_to_weak_absolute_boundary']:.1f} °C", f"|Δ|={thresholds['nanomelt_tm_c']['weak_to_moderate_absolute_boundary']:.1f} °C", f"|Δ|={thresholds['nanomelt_tm_c']['moderate_to_strong_absolute_boundary']:.1f} °C", "Δ>0有利"],
        ],
        [1.0, 1.25, 1.25, 1.25, 0.8],
        size=8.4,
    )
    _paragraph(doc, "这些阈值是预声明的工程幅度档，用来防止极小数值差异驱动选择；它们不是从Nb252突变体实验效应校准得到的。候选在同一幅度档内不按原始小数精确排序。")
    _figure(
        doc,
        Path(data["figures"]["single_landscape"]),
        "图3  允许单突空间的性质景观。该图用于观察全空间分布，不直接等同于最终候选排序；AntiFold只承担负向排除。",
        width_cm=16.2,
    )


def _parent_rows(rows: Sequence[Mapping[str, Any]]) -> list[list[str]]:
    output = []
    for row in rows:
        output.append(
            [
                row["v3_parent_panel_order_not_efficacy_rank"],
                _mutation(row),
                row["region"],
                _metric(row["netsolp_delta_u"], row["netsolp_u_band_v3"]),
                _metric(row["netsolp_delta_s"], row["netsolp_s_band_v3"]),
                _metric(row["nanomelt_delta_tm_c"], row["nanomelt_tm_band_v3"], tm=True),
                f"{_structure_source(row['antifold_selection_source'])}\n{EXPERT_LABELS.get(row['expert_structural_assessment'], row['expert_structural_assessment'])}（{EXPERT_LABELS.get(row['expert_confidence'], row['expert_confidence'])}置信）",
                _parent_note(row),
            ]
        )
    return output


def _add_parent_selection(doc: Document, data: Mapping[str, Any], figures: Mapping[str, Path]) -> None:
    _new_section(doc, "6. 从847条单突到15条父单突")
    facts = data["parent_selection"]["facts"]
    _paragraph(doc, "这里的“硬约束”指禁止改变24个实验直接界面位点、天然保守冻结位点、Cys22/Cys95和末端125–128位SSGS，同时禁止新增Cys；“序列硬风险”指新增Pro造成主链构象约束、突变所在7-aa窗口出现至少6个疏水残基且比WT增加，或任一7-aa窗口的净侧链电荷绝对值达到至少5且比WT增加。单突只有在不触发这些规则和AntiFold双条件排除后，才以U/S/Tm的中等或显著改善作为主要正向证据。")
    _paragraph(doc, "61条性质合格候选经位置与替换多样性抽取形成30条代表性短名单；另加入T99F作为“稳定词”探索项，形成31条专家审查池。稳定词是用简并氨基酸符号表示、被假设可能与稳定性有关的短序列模式，在本项目中仅作为探索性软证据。")
    _paragraph(doc, f"专家审查仅在有高置信度、具体且不可由其他证据合理抵消的物理风险时执行硬排除，共排除{facts['high_confidence_expert_risk_exclusion_count']}条；其余担忧作为降级或实验注释。最终15条覆盖{facts['selected_unique_position_count']}个位点，其中9条至少有一项显著有利指标，5条以多指标中等证据或机制互补入选，T99F为唯一探索例外。")
    _table(
        doc,
        ["排除突变", "主要物理风险", "专家排除依据"],
        [
            ["A49F", "埋藏核心过度包装", "小体积Ala替换为大芳香侧链，实验结构支持高置信空间拥挤风险。"],
            ["A49M", "埋藏核心过度包装与氧化", "埋藏Ala替换为较大Met，兼有空间拥挤和新增氧化敏感性。"],
            ["S50F", "局部过度包装", "部分埋藏且靠近受体的Ser替换为大芳香侧链，局部碰撞与疏水暴露风险同向。"],
            ["R71G", "极性网络与构象稳定性损失", "去除内部Arg极性接触，并在非典型转角位置引入Gly柔性。"],
            ["A96R", "埋藏电荷与二硫键邻近", "深埋Ala替换为大体积带电Arg，且邻近Cys95，包装与折叠风险明确。"],
        ],
        [0.8, 1.45, 3.1],
        size=8.2,
    )
    _paragraph(doc, "T99F没有U/S/Tm中等或显著改善，不应与性质支持候选混称；它因新增稳定词假设被保留为探索对照，且已明确记录部分埋藏、界面邻近及S/Tm轻度不利等风险。")
    _figure(
        doc,
        Path(figures["single_selection_flow"]),
        "图4  单突选择链。AntiFold只减少高风险状态，绝不产生正向候选；展示顺序不是效力排名。",
        width_cm=16.4,
    )
    _figure(
        doc,
        Path(figures["parent15_property_heatmap"]),
        "图5  15条父单突的U/S/Tm幅度档。颜色表示预设幅度类别；AF3补充、潜在化学风险与T99F探索属性仅作注释。",
        width_cm=16.4,
    )

    _new_section(doc, "7. 15条父单突的性质与专家复核")
    _paragraph(doc, "“父单突”是指获准作为双突组成部分的单突。下表按固定展示顺序列出15条父单突；顺序用于交付与追溯，不是1–15的精确效力排名。结构意见同时考虑暴露程度、局部包装、主链构象、极性/电荷变化、二硫键邻近和潜在化学风险。溶解度与热稳定性判断是结合结构机制与预测方向形成的专家假设，不是实测结果。")
    rows = _parent_rows(data["parent_selected15"])
    headers = ["序", "单突", "区域", "ΔU", "ΔS", "ΔTm", "结构证据/判断", "入选依据与风险"]
    weights = [0.38, 0.62, 0.56, 0.82, 0.82, 0.92, 1.45, 2.55]
    _table(doc, headers, rows, weights, size=7.15)
    _paragraph(doc, "表中“实验复合物”指该reported位点具有实验坐标并以复合物环境为主要证据；“AF3补充上下文”表示实验坐标缺失，只能独立参考AF3单体预测，不能写作实验结构结论。")


def _double_rows(rows: Sequence[Mapping[str, Any]]) -> list[list[str]]:
    output = []
    for row in rows:
        output.append(
            [
                row["final_double_panel_order_not_efficacy_rank"],
                row["mutation_set"],
                f"{row['region_a']} + {row['region_b']}",
                _metric(row["netsolp_u_delta_vs_wt"], row["netsolp_u_magnitude_band"]),
                _metric(row["netsolp_s_delta_vs_wt"], row["netsolp_s_magnitude_band"]),
                _metric(row["nanomelt_tm_c_delta_vs_wt"], row["nanomelt_tm_c_magnitude_band"], tm=True),
                f"{_structure_source(row['pair_structure_distance_source'])}\n{_soft_risk(row['effective_soft_sequence_risk_flags'])}",
                _double_note(row),
            ]
        )
    return output


def _add_double_selection(doc: Document, data: Mapping[str, Any], figures: Mapping[str, Path]) -> None:
    _new_section(doc, "8. 102条双突的完整评价")
    facts = data["double_selection"]["facts"]
    _paragraph(doc, f"15条父单突形成105个无序理论配对；L11、F30和K75各有一对同位点替换不能共存，因此去除3对后得到102条有效双突。102条均从完整128-aa序列重建并重新运行NetSolP U/S与NanoMelt Tm，没有用两个单突分数相加代替双突实算。")
    _paragraph(doc, "58条“详细复核”和44条“标准复核”只表示专家记录的详略：至少两项中等/显著改善、存在中等/显著不利，或两位点在WT几何上相邻者进入详细复核。两组使用相同入选规则，复核深度本身既不加分也不淘汰。T99F组合没有强制名额、加分或排除规则，与其他双突平等评价。")
    _paragraph(doc, "AntiFold在双突阶段仍仅是风险排除背景：两个组成单突都必须通过单突双条件排除门；不计算双突AntiFold分数、不相加两个数值，也不把AntiFold改善写成入选理由。")
    _figure(
        doc,
        Path(figures["double_selection_flow"]),
        "图6  双突完整空间与终选过程。102条全部完成U/S/Tm重算；复核深度不等于筛选层级。",
        width_cm=16.4,
    )
    _paragraph(doc, f"最终15条双突中，{facts['selected_three_metric_positive_count']}条在三个指标上均达到中等或显著有利，{facts['selected_two_metric_positive_count']}条在两个指标上达到中等或显著有利；没有中等或显著不利指标。组合覆盖13/15个父单突和10/12个reported位点，15个位置对均不重复，且没有两位点在WT结构中属于局部邻近组合。")

    _new_section(doc, "9. 15条双突的性质与组合依据")
    _figure(
        doc,
        Path(figures["double15_property_heatmap"]),
        "图7  15条入选双突的U/S/Tm幅度档。6条为三指标中等/显著有利，9条为两指标中等/显著有利；展示顺序不是效力排名。",
        width_cm=16.4,
    )
    _paragraph(doc, "双突入选综合考虑性质幅度档、组成单突专家意见、完整序列潜在化学风险、WT位点空间关系和组合多样性。该过程是显式人工复核后的实验面板选择，不声称由唯一优化函数得到全局最优。")
    rows = _double_rows(data["selected_doubles15"])
    headers = ["序", "双突", "区域", "ΔU", "ΔS", "ΔTm", "结构证据/风险", "入选依据"]
    weights = [0.38, 1.05, 0.86, 0.78, 0.78, 0.9, 1.42, 2.33]
    _table(doc, headers, rows, weights, size=7.05)


def _add_final_panel(doc: Document, data: Mapping[str, Any]) -> None:
    _new_section(doc, "10. 最终30条候选面板")
    _paragraph(doc, "最终面板由15条单突和15条双突组成，WT作为面板之外的独立对照。所有候选序列唯一、长度128 aa、末端SSGS保持不变、Cys仍仅位于reported 22与95；没有候选命中冻结的24个直接界面位点。")
    singles = [row for row in data["final_panel30"] if row["candidate_kind"] == "single_mutant"]
    doubles = [row for row in data["final_panel30"] if row["candidate_kind"] == "double_mutant"]
    _heading(doc, "10.1 单突（15条）")
    _table(
        doc,
        ["展示序", "突变", "reported位置", "区域", "实验假设摘要"],
        [
            [row["final_panel_order_not_efficacy_rank"], row["mutation_set"], row["reported_positions_1based"], row["regions"], _parent_note(next(parent for parent in data["parent_selected15"] if _mutation(parent) == row["mutation_set"]))]
            for row in singles
        ],
        [0.55, 0.8, 0.9, 0.75, 3.2],
        size=8.1,
    )
    _heading(doc, "10.2 双突（15条）")
    selected_by_mutation = {row["mutation_set"]: row for row in data["selected_doubles15"]}
    _table(
        doc,
        ["展示序", "组合", "reported位置", "区域", "实验假设摘要"],
        [
            [row["final_panel_order_not_efficacy_rank"], row["mutation_set"], row["reported_positions_1based"], row["regions"], _double_note(selected_by_mutation[row["mutation_set"]])]
            for row in doubles
        ],
        [0.55, 1.1, 1.0, 0.85, 3.4],
        size=7.9,
    )
    _paragraph(doc, "上述顺序用于样品管理和结果回读，不应解释为连续效力排名。性质预测的真实小数与完整序列见附录；实验决策应以WT配对比较为准。")


def _add_risks(doc: Document, data: Mapping[str, Any]) -> None:
    _new_section(doc, "11. 候选级风险与适用边界", page_break_before=False)
    _table(
        doc,
        ["风险/边界", "涉及候选", "解释与处置"],
        [
            ["稳定词探索例外", "T99F", "只有稳定词新增假设；S/Tm轻度不利且有中等结构担忧。作为探索对照，不代表性质预测支持。"],
            ["二硫键邻近", "A23R；A23R+S55G", "A23靠近Cys22，大体积带电Arg可能影响局部包装；当前为中等置信担忧。"],
            ["新增脱酰胺风险", "F30N；S55G+F30N", "新增reported 30的NG基序，可能增加脱酰胺敏感性。"],
            ["氧化敏感性", "L11M；N76G+L11M", "增加一个Met氧化敏感残基，可能提高储存或处理过程中的氧化风险。"],
            ["实验坐标缺失", "L11Y、F29Q、L11M及4条含L11双突", "共7/30个构建依赖AF3补充位置上下文；不能等同于实验复合物证据。"],
            ["CDR1缺口边界", "F30S、F30N及4条含F30双突", "共6/30属于同一F30机制家族；实验结构邻近未解析片段，不能视为6份独立结构证据。"],
            ["AntiFold门边界", "8/15个父单突", "ΔlogP≤−3但未落入位置内最差4名，因此未触发双条件排除；“通过”不等于AntiFold推荐。"],
            ["双突结构建模", "全部15条双突", "没有双侧链重建、骨架松弛或双突AntiFold；WT位点远距只降低直接局部耦联担忧。"],
            ["结合保持", "全部30条", "24个直接界面位点未突变，但非界面/CDR构象变化仍可能间接影响结合；当前尚未验证结合保持。"],
        ],
        [1.2, 1.45, 3.45],
        size=8.25,
    )
    _heading(doc, "预测与审计边界", page_break_before=True)
    _bullet(doc, "NetSolP、NanoMelt和AntiFold输出均为预测，不是Nb252突变体的实测表达量、溶解度、Tm或结构。")
    _bullet(doc, "U与S是NetSolP的两个分别保留输出，不应解释为两个统计独立模型。")
    _bullet(doc, "NanoMelt对WT与候选一致评价前126-aa编号域，并裁去末端GS；最终设计序列仍为完整128 aa。")
    _bullet(doc, "双突预测器输出的非加和残差不等于物理上位性，也未用于最终选择。")
    _bullet(doc, "用于本报告的双突结果矩阵已经过候选身份与序列一致性检查；但远程运行的最原始103行紧凑评分表未随本地报告材料归档，因此无法仅凭当前本地材料从原始输出重新构建该矩阵。")
    _paragraph(doc, "独立审计未发现序列错配、突变重建错误、冻结位点违规或最终CSV/FASTA不一致；总判定为可继续实验准备，但需保留上述科学与溯源限制。详细审计与本报告置于同一V3报告目录。")


def _add_appendix(doc: Document, data: Mapping[str, Any]) -> None:
    _new_section(doc, "附录A  指标和符号说明")
    _table(
        doc,
        ["术语", "报告中的含义"],
        [
            ["reported position", "完整128-aa母本中的1-based位置；不是IMGT编号"],
            ["U / S", "NetSolP综合可用性/预测溶解性两个输出，范围0–1，Δ相对同批WT"],
            ["预测Tm", "NanoMelt预测表观熔解温度；本项目一致评分前126 aa"],
            ["AntiFold通过", "未同时满足ΔlogP≤−3和位置内最差4/20；不表示正向支持"],
            ["实验复合物", "该位点有实验坐标，并以NK2R–Nb252复合物为主结构证据"],
            ["AF3补充上下文", "实验坐标缺失时使用的独立预测结构上下文"],
            ["稳定词", "简并氨基酸模式的探索性序列特征，仅作软证据"],
        ],
        [1.25, 3.9],
        size=8.7,
    )

    _new_section(doc, "附录B  母本与最终30条完整序列")
    _paragraph(doc, "以下均为128-aa完整构建序列。序列目录用于合成与回读；突变标签仍以reported-sequence编号解释。")
    _heading(doc, "B.1 WT母本")
    _table(doc, ["构建", "完整序列（128 aa）"], [["Nb252 WT", data["parent_sequence"]]], [0.9, 4.8], size=9.0, mono_columns={1})
    rows = data["final_panel30"]
    _heading(doc, "B.2 单突序列")
    _table(
        doc,
        ["展示序", "突变", "完整序列（128 aa）"],
        [[row["final_panel_order_not_efficacy_rank"], row["mutation_set"], row["sequence"]] for row in rows if row["candidate_kind"] == "single_mutant"],
        [0.55, 0.9, 5.0],
        size=9.0,
        mono_columns={2},
    )
    _heading(doc, "B.3 双突序列")
    _table(
        doc,
        ["展示序", "组合", "完整序列（128 aa）"],
        [[row["final_panel_order_not_efficacy_rank"], row["mutation_set"], row["sequence"]] for row in rows if row["candidate_kind"] == "double_mutant"],
        [0.55, 1.15, 4.75],
        size=9.0,
        mono_columns={2},
    )


def build_v3_report_document(
    *,
    template_docx: Path,
    output_docx: Path,
    data: Mapping[str, Any],
    report_figures: Mapping[str, Path],
) -> dict[str, Any]:
    """Build the V3 DOCX while preserving the historical report as a template only."""
    if not template_docx.is_file():
        raise V3ReportDocumentError(f"Historical visual template is missing: {template_docx}")
    required_figures = {
        "tool_validation_summary",
        "single_selection_flow",
        "parent15_property_heatmap",
        "double_selection_flow",
        "double15_property_heatmap",
    }
    if not required_figures <= report_figures.keys():
        raise V3ReportDocumentError("V3 report figure set is incomplete")
    doc = Document(str(template_docx))
    _clear_template_body(doc)
    doc.core_properties.title = "Nb252纳米抗体BL21表达量优化项目报告（V3）"
    doc.core_properties.subject = "15条单突和15条双突最终候选面板"
    doc.core_properties.author = "Antibody_optimization project"
    _style_document(doc)
    _set_footer(doc)
    _add_cover(doc, data)
    _add_executive_summary(doc, data)
    _add_parent_and_constraints(doc, data)
    _add_conservation(doc, data, report_figures)
    _add_tools(doc, data, report_figures)
    _add_magnitude(doc, data)
    _add_parent_selection(doc, data, report_figures)
    _add_double_selection(doc, data, report_figures)
    _add_final_panel(doc, data)
    _add_risks(doc, data)
    _add_appendix(doc, data)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    digest = hashlib.sha256(output_docx.read_bytes()).hexdigest()
    return {
        "output_docx": output_docx.as_posix(),
        "sha256": digest,
        "candidate_count": len(data["final_panel30"]),
        "single_count": sum(row["candidate_kind"] == "single_mutant" for row in data["final_panel30"]),
        "double_count": sum(row["candidate_kind"] == "double_mutant" for row in data["final_panel30"]),
        "antifold_role": data["antifold_policy"],
    }


def write_v3_report_manifest(
    *,
    path: Path,
    document_metadata: Mapping[str, Any],
    pdf_path: Path | None,
    template_docx: Path,
    source_artifacts: Mapping[str, str],
    figure_metadata: Mapping[str, Any],
    repository_root: Path,
) -> None:
    """Write compact report provenance after DOCX/PDF generation and QA."""
    root = repository_root.resolve()

    def portable_path(value: str | Path) -> str:
        resolved = Path(value).resolve()
        try:
            relative = Path(os.path.relpath(resolved, start=root))
        except ValueError as exc:
            raise V3ReportDocumentError(
                f"Report manifest path is on a different filesystem: {resolved}"
            ) from exc
        return relative.as_posix()

    document_payload = dict(document_metadata)
    document_payload["output_docx"] = portable_path(document_payload["output_docx"])
    source_payload = {
        name: portable_path(value) for name, value in source_artifacts.items()
    }
    figure_payload: dict[str, Any] = {}
    for name, metadata in figure_metadata.items():
        item = dict(metadata)
        for path_key in ("path", "png", "svg"):
            if path_key in item:
                item[path_key] = portable_path(item[path_key])
        figure_payload[name] = item

    payload: dict[str, Any] = {
        "schema_version": 1,
        "report": "Nb252_BL21_expression_optimization_V3",
        "status": "generated_pending_or_completed_visual_QA",
        "historical_template_sha256": hashlib.sha256(template_docx.read_bytes()).hexdigest(),
        "document": document_payload,
        "source_artifacts": source_payload,
        "figures": figure_payload,
        "scope": {
            "ppt_created": False,
            "delivery_archive_created": False,
            "historical_v2_assets_modified": False,
            "antifold_role": "negative_risk_exclusion_only",
        },
    }
    if pdf_path is not None:
        if not pdf_path.is_file():
            raise V3ReportDocumentError(f"Requested PDF manifest binding is missing: {pdf_path}")
        payload["pdf"] = {
            "path": portable_path(pdf_path),
            "sha256": hashlib.sha256(pdf_path.read_bytes()).hexdigest(),
        }
        payload["status"] = "generated_and_bound"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
